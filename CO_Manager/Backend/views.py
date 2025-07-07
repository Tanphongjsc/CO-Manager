from django.shortcuts import render, get_object_or_404, redirect
from django.views.decorators.http import require_GET, require_POST, require_http_methods
from django.views.decorators.csrf import csrf_exempt
from django.http import JsonResponse, HttpResponse
from django.utils import timezone
from django.db import transaction
from django.forms.models import model_to_dict

from .models import *
from django.db.models import Q

from datetime import datetime, date
from collections import defaultdict
from io import BytesIO
import json
import requests
import re
import os

from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from openpyxl.utils import get_column_letter

from .pdf_utils import convert_number_to_vietnamese_words, generate_excel
from collections import defaultdict
import xlsxwriter
from io import BytesIO
from num2words import num2words
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
# Các view đã có
def dashboard(request):
    return render(request, 'dashboard.html', {})

def orders(request):
    return render(request, 'orders.html', {})

def ctc_ledger(request):
    """Trang quản lý CTC (Chi tiết công việc)"""
    
    # Lấy danh sách CTC với select_related để giảm query
    ctc = BangKeCtc.objects.select_related('id_lenh_san_xuat', 'id_san_pham').order_by('id_bang_ke_ctc')
    

    # Lấy chi tiết lệnh sản xuất với prefetch_related để tối ưu query
    ct_lenh_sx = CtLenhSanXuat.objects.all().select_related('id_lenh_san_xuat', 'id_san_pham', 'id_nguyen_vat_lieu')

    # Lấy danh sách lệnh sản xuất duy nhất và sắp xếp theo thứ tự
    lenh_san_xuat_options = dict(enumerate(
        sorted(set(item.id_lenh_san_xuat.id_lenh_san_xuat for item in ct_lenh_sx), reverse=True)
    ))

    orders = {}
    
    # Xây dựng orders dictionary với logic tối ưu
    for item in ct_lenh_sx:
        lenh_id = item.id_lenh_san_xuat.id_lenh_san_xuat
        san_pham_id = item.id_san_pham.id_san_pham
        
        # Sử dụng setdefault để khởi tạo nested dictionary một cách ngắn gọn
        if lenh_id not in orders:
            orders[lenh_id] = {}
            
        if san_pham_id not in orders[lenh_id]:
            orders[lenh_id][san_pham_id] = {
                "san_pham": model_to_dict(item.id_san_pham),
                "so_luong_san_pham": item.so_luong_san_pham,
                "nguyen_vat_lieu": {}
            }
        
        # Thêm nguyên vat liệu nếu tồn tại
        if item.id_nguyen_vat_lieu:
            nvl_id = item.id_nguyen_vat_lieu.id_san_pham
            dict_nvl = model_to_dict(item.id_nguyen_vat_lieu)
            dict_nvl["so_luong_nguyen_vat_lieu"] = item.so_luong_nguyen_vat_lieu
            orders[lenh_id][san_pham_id]["nguyen_vat_lieu"][nvl_id] = dict_nvl
    
    # Chuyển đổi sang list format bằng dictionary comprehension
    orders_list = {
        lenh_id: list(products.values()) 
        for lenh_id, products in orders.items()
    }
    
    context = {
        'ctc_list': ctc,
        "lenh_san_xuat_options": lenh_san_xuat_options,
        "orders": json.dumps(orders_list, ensure_ascii=False)  # Thêm ensure_ascii=False cho tiếng Việt
    }
    
    return render(request, 'ctc_management.html', context)



def ctc_detail(request, pk):
    """Trang chi tiết CTC"""

    chi_tiet_ctc = CtBangKeCtc.objects.filter(id_bang_ke_ctc=pk).select_related('id_bang_ke_ctc','id_bang_ke_ctc__id_san_pham')
    if not chi_tiet_ctc.exists():
        return JsonResponse({'success': False, 'message': 'Không tìm thấy dữ liệu'}, status=404)
    
    # Lấy thông tin bảng kê chính từ record đầu tiên
    ctc = chi_tiet_ctc[0].id_bang_ke_ctc
    
    # Lấy thông tin nguyên vật liệu một lần
    id_nvl = [item.id_san_pham for item in chi_tiet_ctc]
    nguyen_vat_lieu = {
        item.id_san_pham: item 
        for item in VatTu.objects.filter(id_san_pham__in=id_nvl).only('id_san_pham', 'ma_hs')
    }
    
    # Chuẩn bị dữ liệu response
    ctc_data = {
        'id_bang_ke_ctc': ctc.id_bang_ke_ctc,
        'id_lenh_san_xuat_id': ctc.id_lenh_san_xuat_id,
        'id_san_pham_actual_id': ctc.id_san_pham.id_san_pham if ctc.id_san_pham else None,
        'id_san_pham': {
            'id_san_pham': ctc.id_san_pham.id_san_pham,
            'ten_khac': ctc.id_san_pham.ten_khac,
            'ma_hs': ctc.id_san_pham.ma_hs,
            'don_vi_tinh': ctc.id_san_pham.don_vi_tinh
        } if ctc.id_san_pham else None,
        'so_to_hai_quan': ctc.so_to_hai_quan,
        'so_luong': ctc.so_luong,
        'tri_gia_fob': ctc.tri_gia_fob,
        'chi_tiet_nguyen_lieu': [
            {
                'id': ct.id_ct_bang_ke_ctc,
                'ten_nguyen_lieu': ct.ten_nguyen_lieu,
                'ma_hs': nguyen_vat_lieu.get(ct.id_san_pham, {}).ma_hs,
                'don_vi_tinh': nguyen_vat_lieu.get(ct.id_san_pham, {}).don_vi_tinh,
                'don_gia': ct.don_gia,
                'dinh_muc_san_pham_hao_hut': ct.dinh_muc_san_pham_hao_hut,
                'thanh_tien_co_xuat_xu_field': ct.thanh_tien_co_xuat_xu_field,
                'thanh_tien_khong_xuat_xu_field': ct.thanh_tien_khong_xuat_xu_field,
                'nuoc_xuat_xu': ct.nuoc_xuat_xu,
                'ngay_ke_bang_thu_mua': ct.ngay_ke_bang_thu_mua,
                'so_ban_khai_bao': ct.so_ban_khai_bao,
                'ngay_bang_ke_wo': ct.ngay_bang_ke_wo,
                'ghi_chu': ct.ghi_chu,
            }
            for ct in chi_tiet_ctc
        ]
    }
    
    if request == None:
        return ctc_data
    
    return JsonResponse({'success': True, 'ctc_data': ctc_data})

@require_POST
def ctc_create(request):
    """Tạo mới bảng kê CTC"""
    try:
        data = json.loads(request.body)
        
        # Lấy các đối tương lệnh sản xuất, sản phẩm tương ứng
        lenh_san_xuat = get_object_or_404(LenhSanXuat, id_lenh_san_xuat=data['id_lenh_san_xuat_id'])
        san_pham = get_object_or_404(VatTu, id_san_pham=data['id_san_pham_id'])
        
        # Tạo mới bảng kê CTC
        with transaction.atomic():
            ctc_object = BangKeCtc.objects.create(
                id_lenh_san_xuat=lenh_san_xuat,
                id_san_pham=san_pham,
                so_to_hai_quan=data['so_to_hai_quan'],
                so_luong=data['so_luong'],
                tri_gia_fob=data['tri_gia_fob']
            )
            
            # Tạo mới hoàng loạt Chi tiết bảng kê CTC
            chi_tiet_nguyen_lieu = data.get('chi_tiet_nguyen_lieu', [])
            if chi_tiet_nguyen_lieu:
                ct_ctc_objs = [
                    CtBangKeCtc(
                        id_bang_ke_ctc=ctc_object,
                        id_san_pham=item['id_nguyen_lieu'],
                        ten_nguyen_lieu=item['ten_nguyen_lieu'],
                        don_gia=item['don_gia'],
                        dinh_muc_san_pham_hao_hut=item['dinh_muc_san_pham_hao_hut'],
                        thanh_tien_co_xuat_xu_field=item['thanh_tien_co_xuat_xu_field'],
                        thanh_tien_khong_xuat_xu_field=item['thanh_tien_khong_xuat_xu_field'],
                        nuoc_xuat_xu=item['nuoc_xuat_xu'],
                        ngay_ke_bang_thu_mua=item['ngay_ke_bang_thu_mua'],
                        so_ban_khai_bao=item['so_ban_khai_bao'],
                        ngay_bang_ke_wo= item['ngay_bang_ke_wo'] if item['ngay_bang_ke_wo'] else None,
                        ghi_chu=item['ghi_chu']
                    )
                    for item in chi_tiet_nguyen_lieu
                ]
                CtBangKeCtc.objects.bulk_create(ct_ctc_objs)
        
        return JsonResponse({
            'success': True,
            'message': 'Tạo mới chứng từ CTC thành công!',
        })
        
    except Exception as e:
        return JsonResponse({'success': False, 'message': f'Invalid JSON data, {e}'}, status=400)


@transaction.atomic
def ctc_update(request, pk):
    """Cập nhật bảng kê CTC với logic rõ ràng, dễ hiểu."""
    try:
        data_update = json.loads(request.body)
        
        # === CẬP NHẬT THÔNG TIN CHÍNH ===
        ctc = get_object_or_404(BangKeCtc, id_bang_ke_ctc=pk)
        
        fields_to_update_main = [
            'id_lenh_san_xuat_id', 'id_san_pham_id', 'so_to_hai_quan', 
            'so_luong', 'tri_gia_fob'
        ]
        for field in fields_to_update_main:
            if field in data_update:
                setattr(ctc, field, data_update[field]) # setattr(object, name, value) tương đương với object.name = value
        ctc.save()

        # === XỬ LÝ DỮ LIỆU CHI TIẾT ===
        chi_tiet_items_request = data_update.get('chi_tiet_nguyen_lieu', [])
        
        existing_details_map = {          # Lấy tất cả các bản ghi chi tiết hiện có trong CSDL để so sánh
            item.id_ct_bang_ke_ctc: item 
            for item in CtBangKeCtc.objects.filter(id_bang_ke_ctc=ctc)
        }

        # Chuẩn bị các danh sách để thực hiện thao tác hàng loạt
        items_to_create = []
        items_to_update = []

        incoming_ids = set()         # Dùng `set` để lưu các ID được gửi lên, giúp xác định các ID cần xóa sau này

        # Hàm nội bộ giúp chuyển đổi chuỗi ngày tháng an toàn
        def parse_date_safely(date_string):
            if isinstance(date_string, str) and date_string:
                return date.fromisoformat(date_string.split('T')[0])
            return None

        # Phân loại từng mục trong dữ liệu gửi lên vào danh sách Cập nhật hoặc Tạo mới
        for item_data in chi_tiet_items_request:
            item_id = item_data.get('id')
            
            # Chuẩn bị một dictionary chứa dữ liệu đã được làm sạch và ánh xạ đúng tên trường model
            prepared_data = {
                'id_san_pham': item_data.get('id_nguyen_lieu'),
                'ten_nguyen_lieu': item_data.get('ten_nguyen_lieu'),
                'don_gia': item_data.get('don_gia'),
                'dinh_muc_san_pham_hao_hut': item_data.get('dinh_muc_san_pham_hao_hut'),
                'thanh_tien_co_xuat_xu_field': item_data.get('thanh_tien_co_xuat_xu_field'),
                'thanh_tien_khong_xuat_xu_field': item_data.get('thanh_tien_khong_xuat_xu_field'),
                'nuoc_xuat_xu': item_data.get('nuoc_xuat_xu'),
                'ngay_ke_bang_thu_mua': item_data.get('ngay_ke_bang_thu_mua'),
                'so_ban_khai_bao': item_data.get('so_ban_khai_bao'),
                'ngay_bang_ke_wo': parse_date_safely(item_data.get('ngay_bang_ke_wo')),
                'ghi_chu': item_data.get('ghi_chu')
            }
            
            if item_id and item_id in existing_details_map:
                # ----- TH1: CẬP NHẬT -----   Nếu có ID và ID này tồn tại trong CSDL -> Cập nhật
                incoming_ids.add(item_id)
                db_item = existing_details_map[item_id]
                for field, value in prepared_data.items():
                    setattr(db_item, field, value)
                items_to_update.append(db_item)
            elif not item_id:
                # ----- TH2: TẠO MỚI -----    Nếu không có ID -> Tạo mới
                prepared_data['id_bang_ke_ctc'] = ctc
                items_to_create.append(CtBangKeCtc(**prepared_data))
        
        # === XÓA CÁC MỤC KHÔNG CÒN TỒN TẠI ===
        ids_to_delete = set(existing_details_map.keys()) - incoming_ids
        if ids_to_delete:
            CtBangKeCtc.objects.filter(id_ct_bang_ke_ctc__in=ids_to_delete).delete()

        # === THỰC THI THAO TÁC VỚI CSDL ===
        if items_to_update:
            # Tự động lấy danh sách các trường cần cập nhật
            update_fields = list(prepared_data.keys()) 
            CtBangKeCtc.objects.bulk_update(items_to_update, update_fields)
            
        if items_to_create:
            CtBangKeCtc.objects.bulk_create(items_to_create)

        return JsonResponse({'success': True, 'message': 'Cập nhật bảng kê CTC thành công!'})

    except Exception as e:
        return JsonResponse({'success': False, 'message': f'Đã xảy ra lỗi: {e}'}, status=500)


def ctc_delete(request, pk):
    """Xóa bảng kê CTC"""

    if pk:
        ctc = get_object_or_404(BangKeCtc, id_bang_ke_ctc=pk)
        ctc.delete()

        return JsonResponse({
            'success': True,
            'message': 'Xóa bảng kê CTC thành công!'
        })
    
    return JsonResponse({
        'success': False,
        'message': 'Không tìm thấy mã bảng kê CTC!'
    })


def ctc_export(request, pk):
    """Xuất báo cáo tỉ lệ phối trộn theo định dạng PDF hoặc Excel """

    format = request.GET.get('format', '').lower()
    response_data = ctc_detail(request=None, pk=pk)

    if format == 'pdf':
        return render(request, 'form/bang_ke_ctc_pdf.html', {'data': response_data})
    elif format == 'excel':
        return create_bang_ke_ctc_excel_response(response_data)
    else:
        return JsonResponse({'success': False, 'message': 'Format không hợp lệ!'}, status=400)

def create_bang_ke_ctc_excel_response(data):
    """Tạo file Excel báo cáo bảng kê khai hàng hoá xuất khẩu đặt tiêu chí CTC."""
    wb = Workbook()
    ws = wb.active
    
    # Thiết lập styles cho Excel
    styles = {
        'border': Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin')),
        'thick_border': Border(left=Side(style='thick'), right=Side(style='thick'), top=Side(style='thick'), bottom=Side(style='thick')),
        'fill_header': PatternFill(start_color="F9F9F9", end_color="F9F9F9", fill_type="solid"),
        'base_font': Font(name='Times New Roman', size=11),
        'bold_font': Font(name='Times New Roman', size=11, bold=True),
        'title_font': Font(name='Times New Roman', size=14, bold=True),
        'info_font': Font(name='Times New Roman', size=10),
        'center': Alignment(horizontal='center', vertical='center', wrap_text=True),
        'left': Alignment(horizontal='left', vertical='center', wrap_text=True),
        'right': Alignment(horizontal='right', vertical='center')
    }
    
    # ===== PHẦN 1: TIÊU ĐỀ =====
    # Phụ lục V
    cell = ws.cell(row=1, column=1, value="Phụ lục V")
    apply_cell_style(cell, font=styles['bold_font'], align=styles['center'])
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=13)
    
    # Tiêu đề chính
    title = "BẢNG KÊ KHAI HÀNG HOÁ XUẤT KHẨU ĐẶT TIÊU CHÍ \"CTC\""
    cell = ws.cell(row=3, column=1, value=title)
    apply_cell_style(cell, font=styles['title_font'], align=styles['center'])
    ws.merge_cells(start_row=3, start_column=1, end_row=3, end_column=13)
    
    # ===== PHẦN 2: THÔNG TIN THƯƠNG NHÂN (Không có viền và màu nền) =====
    info_start_row = 5
    
    # Thông tin bên trái - các dòng liền kề
    left_info_data = [
        ("Tên Thương Nhân:", "Công ty cổ phần Tân Phong"),
        ("Mã số thuế:", "2600274542"),
        ("Tờ khai Hải quan xuất khẩu số:", data.get('so_to_hai_quan', ''))
    ]
    
    for i, (label, value) in enumerate(left_info_data):
        current_row = info_start_row + i  # Các dòng liền kề
        
        # Label (không có border và background)
        cell = ws.cell(row=current_row, column=1, value=label)
        apply_cell_style(cell, font=styles['bold_font'], align=styles['left'])
        ws.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=2)
        
        # Value (không có border)
        cell = ws.cell(row=current_row, column=3, value=value)
        apply_cell_style(cell, font=styles['base_font'], align=styles['left'])
        ws.merge_cells(start_row=current_row, start_column=3, end_row=current_row, end_column=6)
    
    # Thông tin bên phải - các dòng liền kề
    
    # Xử lý safe attribute access
    ten_khac = data.get('id_san_pham', {}).get('ten_khac', '')
    ma_hs = data.get('id_san_pham', {}).get("ma_hs", '')
    don_vi_tinh = getattr(data.get('id_san_pham', {}), 'don_vi_tinh', '')
    
    right_info_data = [
        ("Tiêu chí áp dụng:", "CC"),
        ("Tên hàng hoá:", ten_khac),
        ("Mã số của hàng hoá (6 số):", ma_hs),
        ("Số lượng:", f"{data.get('so_luong', 0):,.0f} {don_vi_tinh}"),
        ("Trị giá (FOB):", f"{data.get('tri_gia_fob', 0)}")
    ]
    
    for i, (label, value) in enumerate(right_info_data):
        current_row = info_start_row + i  # Các dòng liền kề
        
        # Label (không có border và background)
        cell = ws.cell(row=current_row, column=8, value=label)
        apply_cell_style(cell, font=styles['bold_font'], align=styles['left'])
        ws.merge_cells(start_row=current_row, start_column=8, end_row=current_row, end_column=9)
        
        # Value (không có border)
        cell = ws.cell(row=current_row, column=10, value=value)
        apply_cell_style(cell, font=styles['base_font'], align=styles['left'])
        ws.merge_cells(start_row=current_row, start_column=10, end_row=current_row, end_column=13)
    
    # ===== PHẦN 3: BẢNG DỮ LIỆU (Giữ nguyên viền và màu) =====
    table_start_row = info_start_row + 5  # Giảm khoảng cách
    
    # Header dòng 1: Tiêu đề chính
    headers_main = [
        ("STT", 1, 1, 1, 1),
        ("Tên Nguyên Liệu", 1, 1, 2, 2),
        ("Mã HS", 1, 1, 3, 3),
        ("Đơn vị tính", 1, 1, 4, 4),
        ("Định mức sản phẩm hao hụt", 1, 1, 5, 5),
        ("Nhu cầu nguyên liệu sử dụng cho lô hàng", 1, 1, 6, 8),
        ("Nước xuất xứ", 1, 1, 9, 9),
        ("Tờ khai hải quan nhập khẩu/Hoá đơn trị giá tăng", 1, 1, 10, 11),
        ("C/O ưu đãi nhập khẩu bản khai của nhà sản xuất nhà cung cấp nguyên liệu trong nước", 1, 1, 12, 13)
    ]
    
    for text, start_row_offset, end_row_offset, start_col, end_col in headers_main:
        cell = ws.cell(row=table_start_row, column=start_col, value=text)
        apply_cell_style(cell, font=styles['bold_font'], border=styles['border'], 
                        align=styles['center'], fill=styles['fill_header'])
        if start_col != end_col:
            ws.merge_cells(start_row=table_start_row, start_column=start_col,
                          end_row=table_start_row, end_column=end_col)
    
    # Header dòng 2: Sub headers
    sub_headers = [
        ("Đơn giá", 6),
        ("Có xuất xứ", 7),
        ("Không có xuất xứ", 8),
        ("Số", 10),
        ("Ngày", 11),
        ("Số", 12),
        ("Ngày", 13)
    ]
    
    for text, col in sub_headers:
        cell = ws.cell(row=table_start_row + 1, column=col, value=text)
        apply_cell_style(cell, font=styles['bold_font'], border=styles['border'], 
                        align=styles['center'], fill=styles['fill_header'])
    
    # Header dòng 3: Số thứ tự cột
    column_numbers = ["(1)", "(2)", "(3)", "(4)", "(5)", "(6)", "(7)", "(8)", "(9)", "(10)", "(11)", "(12)", "(13)"]
    for i, num in enumerate(column_numbers, 1):
        cell = ws.cell(row=table_start_row + 2, column=i, value=num)
        apply_cell_style(cell, font=styles['info_font'], border=styles['border'], 
                        align=styles['center'], fill=styles['fill_header'])
    
    # Merge cells cho các cột không có sub-header
    merge_ranges = [
        (table_start_row, table_start_row + 1, 1, 1),  # STT
        (table_start_row, table_start_row + 1, 2, 2),  # Tên Nguyên Liệu
        (table_start_row, table_start_row + 1, 3, 3),  # Mã HS
        (table_start_row, table_start_row + 1, 4, 4),  # Đơn vị tính
        (table_start_row, table_start_row + 1, 5, 5),  # Định mức
        (table_start_row, table_start_row + 1, 9, 9),  # Nước xuất xứ
    ]
    
    for start_row, end_row, start_col, end_col in merge_ranges:
        ws.merge_cells(start_row=start_row, start_column=start_col,
                      end_row=end_row, end_column=end_col)
    
    # Tăng chiều cao cho header
    for row in range(table_start_row, table_start_row + 3):
        ws.row_dimensions[row].height = 30
    
    # ===== PHẦN 4: DỮ LIỆU NGUYÊN LIỆU =====
    data_start_row = table_start_row + 3
    chi_tiet_nguyen_lieu = data.get('chi_tiet_nguyen_lieu', [])
    
    for row_idx, material in enumerate(chi_tiet_nguyen_lieu):
        current_row = data_start_row + row_idx
        
        row_data = [
            (1, row_idx + 1, '#,##0'),
            (2, material.get('ten_nguyen_lieu', ''), None),
            (3, material.get('ma_hs', ''), None),
            (4, material.get('don_vi_tinh', ''), None),
            (5, material.get('dinh_muc_san_pham_hao_hut', 0), '#,##0'),
            (6, material.get('don_gia', 0), '#,##0'),
            (7, material.get('thanh_tien_co_xuat_xu_field', 0), '#,##0'),
            (8, material.get('thanh_tien_khong_xuat_xu_field', 0), '#,##0'),
            (9, material.get('nuoc_xuat_xu', ''), None),
            (10, "Bảng kê thu mua", None),
            (11, material.get('ngay_ke_bang_thu_mua', ''), None),
            (12, material.get('so_ban_khai_bao', ''), None),
            (13, '', None)
        ]
        
        # Xử lý ngày bản khai
        ngay_bang_ke_wo = material.get('ngay_bang_ke_wo', '')
        if ngay_bang_ke_wo:
            try:
                if hasattr(ngay_bang_ke_wo, 'strftime'):
                    row_data[12] = (13, ngay_bang_ke_wo.strftime('%d/%m/%Y'), None)
                else:
                    row_data[12] = (13, str(ngay_bang_ke_wo), None)
            except:
                row_data[12] = (13, '', None)
        
        # Áp dụng dữ liệu và styling (giữ nguyên border cho bảng)
        for col, value, number_format in row_data:
            cell = ws.cell(row=current_row, column=col, value=value)
            
            # Căn lề và styling
            if col == 2:  # Tên nguyên liệu
                align = styles['left']
            elif col in [5, 6, 7, 8]:  # Các cột số liệu
                align = styles['right']
            else:
                align = styles['center']
            
            apply_cell_style(cell, font=styles['base_font'], border=styles['border'], 
                           align=align, number_format=number_format)
    
    # ===== PHẦN 5: KẾT LUẬN VÀ CAM KẾT (Không có viền) =====
    footer_row = data_start_row + len(chi_tiet_nguyen_lieu)  # Giảm khoảng cách
    
    # Kết luận (không có border và background)
    cell = ws.cell(row=footer_row, column=1, value="Kết luận: hàng hoá đáp ứng tiêu chí \"CC\"")
    apply_cell_style(cell, font=styles['bold_font'], align=styles['left'])
    ws.merge_cells(start_row=footer_row, start_column=1, end_row=footer_row, end_column=13)
    
    # Cam kết (liền kề, không có border)
    commitment = "Công ty cam kết số liệu, thông tin khai báo trên là đúng và chịu trách nhiệm trước pháp luật về thông tin, số liệu đã khai."
    cell = ws.cell(row=footer_row + 1, column=1, value=commitment)
    apply_cell_style(cell, font=styles['base_font'], align=styles['left'])
    ws.merge_cells(start_row=footer_row + 1, start_column=1, end_row=footer_row + 1, end_column=13)
    
    # ===== PHẦN 6: CHỮ KÝ (Không có viền) =====
    sig_row = footer_row + 2  # Giảm khoảng cách
    
    # Xử lý ngày tháng
    try:
        current_date = data.get('current_date')
        if current_date is None:
            from datetime import datetime
            current_date = datetime.now()
        elif isinstance(current_date, str):
            from datetime import datetime
            current_date = datetime.strptime(current_date, '%Y-%m-%d')
    except:
        from datetime import datetime
        current_date = datetime.now()
    
    # Ngày tháng (không có border)
    date_str = f"Ngày {current_date.day} tháng {current_date.month} năm {current_date.year}"
    cell = ws.cell(row=sig_row, column=9, value=date_str)
    apply_cell_style(cell, font=styles['base_font'], align=styles['center'])
    ws.merge_cells(start_row=sig_row, start_column=9, end_row=sig_row, end_column=13)
    
    # Người đại diện (liền kề, không có border)
    cell = ws.cell(row=sig_row + 1, column=9, value="NGƯỜI ĐẠI DIỆN THEO PHÁP LUẬT")
    apply_cell_style(cell, font=styles['bold_font'], align=styles['center'])
    ws.merge_cells(start_row=sig_row + 1, start_column=9, end_row=sig_row + 1, end_column=13)
    
    # Hướng dẫn ký (cách 3 dòng để có chỗ ký)
    ws.merge_cells(start_row=sig_row + 2, start_column=9, end_row=sig_row + 2, end_column=13)
    cell = ws.cell(row=sig_row + 2, column=9, value="(Ký, đóng dấu, ghi rõ họ tên)")
    apply_cell_style(cell, font=styles['base_font'], align=styles['center'])
    
    # ===== PHẦN 7: GHI CHÚ (Không có viền) =====
    note_row = sig_row + 7  # Giảm khoảng cách
    
    # Tiêu đề ghi chú
    ws.merge_cells(start_row=note_row, start_column=1, end_row=note_row, end_column=2)     # Hợp nhất ô từ cột 1 đến cột 2 tại dòng note_row
    cell = ws.cell(row=note_row, column=1, value="Ghi chú:")     # Gán giá trị cho ô đầu tiên của vùng hợp nhất
    apply_cell_style(cell, font=styles['bold_font'], align=styles['left'])
    
    # Nội dung ghi chú (liền kề)
    note_text = ("Thương nhân nộp bản in Tờ khai hải quan nhập khẩu và bản sao các chứng từ "
                "(có dấu sao y bản chính): Hoá đơn trị giá tăng, C/O ưu đãi nhập khẩu bản khai "
                "của nhà sản xuất/ nhà cung cấp nguyên liệu trong nước để đối chiếu với thông tin "
                "kê khai từ cột (9) đến cột (13)")
    cell = ws.cell(row=note_row + 1, column=1, value=note_text)
    apply_cell_style(cell, font=styles['info_font'], align=styles['left'])
    ws.merge_cells(start_row=note_row + 1, start_column=1, end_row=note_row + 1, end_column=13)
    
    # ===== PHẦN 8: ĐIỀU CHỈNH CHIỀU RỘNG CỘT =====
    # Điều chỉnh chiều rộng cột để cân đối hơn
    for column in ws.columns:
        max_length = 0
        column_letter = get_column_letter(column[0].column)
        
        for cell in column[10:12+len(chi_tiet_nguyen_lieu)]:  # Chỉ kiểm tra các dòng header và một vài dòng đầu tiên
            try:
                cell_length = len(str(cell.value)) if cell.value else 0
                if cell_length > max_length:
                    max_length = cell_length
            except:
                pass
        
        ws.column_dimensions[column_letter].width = max_length + 3
    
    ws.row_dimensions[note_row + 1].height = 35
    
    # ===== TẠO RESPONSE =====
    output = BytesIO()
    wb.save(output)
    output.seek(0)
    
    # Tạo tên file
    so_to_hai_quan = data.get('so_to_hai_quan', 'Unknown')
    today = current_date.strftime('%Y-%m-%d')
    filename = f"Bang_Ke_CTC_{so_to_hai_quan}_{today}.xlsx"
    
    response = HttpResponse(
        output.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


def get_data_for_ctc_create(request):
    """Lấy dữ liệu fill vào bảng khi tạo mới bảng kê ctc"""
    
    data = json.loads(request.body)
    id_lenh_san_xuat = data.get('id_lenh_san_xuat')

    if id_lenh_san_xuat:
        bang_ke_thu_mua = BangKeThuMuaTuDan.objects.filter(id_lenh_san_xuat = id_lenh_san_xuat).values()
        bang_ke_wo = BangKeWo.objects.filter(id_lenh_san_xuat = id_lenh_san_xuat).values()
        nguyen_vat_lieu = VatTu.objects.filter(
            id_san_pham__in=[item['id_san_pham'] for item in bang_ke_thu_mua] + 
                            [item['id_san_pham'] for item in bang_ke_wo]
        ).values('id_san_pham', 'ten_khac', 'ma_hs', 'don_vi_tinh')

        # Chuyển đổi nguyen_vat_lieu thành dictionary để dễ truy cập
        nguyen_vat_lieu = {item['id_san_pham']: item for item in nguyen_vat_lieu}

        data = {
            "bang_ke_thu_mua": {nguyen_vat_lieu[item['id_san_pham']]['ten_khac']: item['ngay_lap_giay_to'] for item in list(bang_ke_thu_mua)},
            "bang_ke_wo": {nguyen_vat_lieu[item['id_san_pham']]['ten_khac']: item['ngay'] for item in list(bang_ke_wo)}
        }

        return JsonResponse({
            'success': True,
            'data': data,
        })
    
    return JsonResponse({
        'success': False,
        'message': 'Không tìm thấy mã lệnh sản xuất!'
    }, status=404)
    

def users_management(request):
    """API endpoint để lấy danh sách người dùng."""
    users = Nguoi.objects.all().order_by('id')

    context = {
        'users': users,
    }
    return render(request, 'user_management.html', context)


@require_POST
def users_create(request):
    """API endpoint để tạo người dùng mới."""
    try:
        data = json.loads(request.body)

        # Chuẩn hóa dữ liệu Vai trò
        if data.get('role') == "buyer":
            data['role'] = 'Người mua'
        elif data.get('role') == "seller":
            data['role'] = 'Người bán'
        else:
            data['role'] = None

        user = Nguoi.objects.create(
            ten=data.get('name'),
            so_cmnd_cccd=data.get('cmnd'),
            ngay_cap_cmnd_cccd=data.get('date'),
            dia_chi=data.get('address'),
            vai_tro=data.get('role'),
        )
        return JsonResponse({
            'success': True,
            'message': 'Tạo người dùng thành công!'
        })
    
    except:
        return JsonResponse({
            'success': False,
            'message': 'Dữ liệu không hợp lệ!'
        })

@require_POST
def users_update(request):
    """API endpoint để cập nhật thông tin người dùng."""
    try:
        data = json.loads(request.body)
        user_id = data.get('id')

        user = get_object_or_404(Nguoi, id=user_id)

        user.ten = data.get('name', user.ten)
        user.so_cmnd_cccd = data.get('cmnd', user.so_cmnd_cccd)
        user.ngay_cap_cmnd_cccd = data.get('date', user.ngay_cap_cmnd_cccd)
        user.dia_chi = data.get('address', user.dia_chi)

        # Xử lý với trường Vai Trò
        if data.get('role') == "buyer":
            user.vai_tro = 'Người mua'
        elif data.get('role') == "seller":
            user.vai_tro = 'Người bán'
        else:
            user.vai_tro = None

        user.save()

        return JsonResponse({
            'success': True,
            'message': 'Cập nhật người dùng thành công!'
        })
    
    except:
        return JsonResponse({
            'success': False,
            'message': 'Dữ liệu không hợp lệ!'
        }, status=400)


@require_POST
def users_delete(request):
    """API endpoint để xóa người dùng."""
    data = json.loads(request.body)

    user_id = data.get('id')
    if not user_id:
        return JsonResponse({
            'success': False,
            'message': 'Không tìm thấy mã người dùng!'
        }, status=404)
    
    user = get_object_or_404(Nguoi, id=user_id)
    user.delete()

    return JsonResponse({
        'success': True,
        'message': 'Xóa người dùng thành công!'
    })


def product_management(request):
    products = VatTu.objects.all().order_by('id_san_pham')
    context = {
        'products': products,
    }
    return render(request, 'product_management.html', context)

@require_POST
def product_update(request):
    """API endpoint để xử lý cập nhật thông tin sản phẩm"""
    try:
        # Phân tích JSON từ request body
        data = json.loads(request.body)
        data['alt_name'] = data.get('alt_name', '').strip().capitalize()
        
        # Lấy ID sản phẩm từ dữ liệu gửi lên
        product_id = data.get('id')

        # Tìm sản phẩm cần cập nhật trong database
        product = get_object_or_404(VatTu, id_san_pham=product_id)
        
        # Cập nhật thông tin sản phẩm
        product.ten_khac = data.get('alt_name', product.ten_khac)
        product.ma_hs = data.get('hs_code', product.ma_hs)
        product.ghi_chu = data.get('note', product.ghi_chu)

        # Xử lý trường số - tỉ lệ thu hồi
        recovery_rate = data.get('recovery_rate')
        if recovery_rate:
            try:
                product.ty_le_thu_hoi = float(recovery_rate)
            except (ValueError, TypeError):
                # Giữ nguyên giá trị cũ nếu dữ liệu không hợp lệ
                pass
        
        # Lưu thay đổi vào database
        product.save()
        
        # Trả về kết quả thành công
        return JsonResponse({
            'success': True,
            'message': 'Cập nhật sản phẩm thành công!'
        })
        
    except Exception as e:
        return JsonResponse({
            'success': False,
            'message': f'Lỗi: {str(e)}'
        }, status=500)

@require_POST
def product_delete(request):
    """Để xử lý xóa sản phẩm"""
    product_id = json.loads(request.body).get('id')

    if not product_id:
        return JsonResponse({
            'success': False,
            'message': 'Không tìm thấy mã sản phẩm!'
        }, status=404)

    product = get_object_or_404(VatTu, id_san_pham=product_id)
    product.delete()

    return JsonResponse({
        'success': True,
        'message': 'Xóa sản phẩm thành công!'
    })


@require_POST
def products_sync_cloudify(request):
    """Xử lý đồng bộ hóa dữ liệu sản phẩm từ Cloudify."""
    try:
        data_cloudify, session = fetch_erp_data(model='danh.muc.vat.tu.hang.hoa', endpoint='search_read', fields=["MA","TEN","DONG_SAN_PHAM","TINH_CHAT","LIST_NHOM_VTHH","DVT_CHINH_ID"])

    except requests.RequestException as e:
        return JsonResponse({
            'success': False,
            'message': f'Lỗi kết nối đến Cloudify: {str(e)}'
        }, status=500)
        
    if data_cloudify:
        adapter_data = {
            "0":'Vật tư hàng hóa', 
            "1":'Thành phẩm', 
            "2":'Dich vụ', 
            "3":'Chỉ là diễn giải', 
            "4":'Bán thành phẩm'
        }
        for item in data_cloudify['records']:
            item['TINH_CHAT'] = adapter_data.get(item['TINH_CHAT'], None)
            item['DVT_CHINH_ID'] = item['DVT_CHINH_ID'][1] if item['DVT_CHINH_ID'] else None

        if data_cloudify['length'] == len(data_cloudify['records']):

            # Lấy ra danh sách mã sản phẩm từ Cloudify
            cloudify_product = {item['MA']: item for item in data_cloudify['records']}

            # Lấy ra danh sách sản phẩm từ database
            db_products = VatTu.objects.filter(id_san_pham__in=cloudify_product.keys())
            
            # Lấy danh sách các mã đã tìm thấy trong database
            found_product_id = {item.id_san_pham for item in db_products}

            # Tạo danh sách các mã sản phẩm không tìm thấy trong database
            not_found_product_id = set(cloudify_product.keys()) - found_product_id

            # Update lại những sản phẩm đã tìm thấy
            for item in db_products:
                item.ten_sp_chinh = cloudify_product[item.id_san_pham].get('TEN')
                item.don_vi_tinh = cloudify_product[item.id_san_pham].get('DVT_CHINH_ID')
                item.loai_sp = cloudify_product[item.id_san_pham].get('TINH_CHAT')
                item.nhom_vthh = cloudify_product[item.id_san_pham].get('LIST_NHOM_VTHH')
        
            # Thêm mới những sản phẩm không tìm thấy
            objs = []
            for id in not_found_product_id:
                objs.append(
                    VatTu(
                        id_san_pham=id,
                        ten_sp_chinh=cloudify_product[id].get('TEN'),
                        ten_khac=cloudify_product[id].get('TEN'),
                        don_vi_tinh=cloudify_product[id].get('DVT_CHINH_ID'),
                        loai_sp=cloudify_product[id].get('TINH_CHAT'),
                        nhom_vthh=cloudify_product[id].get('LIST_NHOM_VTHH')
                    )
                )

            # Dùng transaction để đảm bảo atomicity
            with transaction.atomic():
                VatTu.objects.bulk_update(db_products, ['ten_sp_chinh', 'don_vi_tinh', 'loai_sp', 'nhom_vthh']) # Cập nhật những sản phẩm đã tìm thấy
                VatTu.objects.bulk_create(objs) # Thêm mới những sản phẩm không tìm thấy

            return JsonResponse({
                'success': True,
                'data': data_cloudify
            })
        
        else:
            return JsonResponse({
                'success': False,
                'message': 'Đồng bộ thiếu dữ liệu!',
                'data': data_cloudify
            })

    return JsonResponse({
        'success': False,
        'message': 'Không thể đồng bộ hóa dữ liệu từ Cloudify!'
    }, status=500)


def fetch_erp_data(model=None, endpoint=None, params=None, fields=None, session=None):
    """Hàm lấy dữ liệu từ hệ thống ERP qua API"""

    # Cấu hình API
    BASE_URL = 'https://tanphongjsc.cloudify.vn/web'
    LOGIN_URL = f'{BASE_URL}/login'
    RPC_URL = f'{BASE_URL}/dataset/{endpoint}'
    EMAIL = os.getenv('cloudify_user')
    PASSWORD = os.getenv('cloudify_password')

    # Tạo session mới nếu chưa có
    if not session:
        session = requests.Session()
        login_page = session.get(LOGIN_URL)
        csrf_token = re.search(r'name="csrf_token".*?value="([^"]+)"', login_page.text).group(1)

        # Đăng nhập
        session.post(LOGIN_URL, data={
            'csrf_token': csrf_token,
            'login': EMAIL,
            'password': PASSWORD,
        })

    # Gọi JSON-RPC
    payload = {
        "jsonrpc": "2.0",
        "method": "call",
        "params": {
            "model": model,
            "fields": fields or [],
            "domain": []
        } if params is None else params,
        "id": int(timezone.now().timestamp() * 1000)
    }

    headers = {
        'Content-Type': 'application/json',
        'X-Requested-With': 'XMLHttpRequest',
    }

    response = session.post(RPC_URL, headers=headers, json=payload)
    data = response.json().get('result', [])
    
    return data, session


def orders(request):
    # Truy vấn danh sách đơn hàng từ cơ sở dữ liệu
    orders = LenhSanXuat.objects.all().order_by('-id_lenh_san_xuat')
    
    for order in orders:
        if order.id_don_hang:
            len_id = len(order.id_don_hang)
            order.id_don_hang = f"ĐH{'0'*(6-len_id)}{int(order.id_don_hang)-1}"
    
    context = {
        'orders': orders,
    }

    return render(request, 'orders.html', context)

def orders_detail(request, pk):
    # Truy vấn chi tiết đơn hàng từ cơ sở dữ liệu
    order_items = CtLenhSanXuat.objects.filter(id_lenh_san_xuat=pk).select_related('id_san_pham', 'id_nguyen_vat_lieu', 'id_lenh_san_xuat')
    if not order_items.exists():
        return render(request, '404.html')

    # Lấy ra danh sách id_san_pham unique
    unique_products = {}
    for item in order_items:
        id = item.id_san_pham
        if id not in unique_products:
            unique_products[id] = item
    
    order_unique_items = list(unique_products.values())        


    # Tính tổng số lượng sản phẩm
    total_quantity = sum(item.so_luong_san_pham for item in order_unique_items)

    context = {
        'orders': order_items,
        'order_unique_items': order_unique_items,
        'total_quantity': total_quantity,
    }

    return render(request, 'orders_detail.html', context)


def fetch_cloudify_orders_data():
    """Lấy dữ liệu từ Cloudify và trả về data đã xử lý"""
    # Lấy danh sách lệnh sản xuất
    production_orders_data, session = fetch_erp_data(model='stock.ex.lenh.san.xuat', endpoint='search_read', fields=['SO_LENH', 'LAP_TU_DON_DAT_HANG_IDS', 'NGAY', 'STOCK_EX_LENH_SAN_XUAT_CHI_TIET_THANH_PHAM_IDS'])

    if not production_orders_data:
        return None, None, None

    production_orders = {order['SO_LENH']: order for order in production_orders_data['records']} # Các lệnh sản xuất, tương ứng dữ liệu cho bảng Lệnh Sản Xuất
    product_norms = {} # Thông tin các sản phẩm của các lệnh sản xuất, tương ứng dữ liệu cho bảng Chi Tiết Lệnh Sản Xuất
    all_detail_ids = [] # ID của các dòng của Chi Tiết Lệnh Sản Xuất

    for order_id, order_data in production_orders.items():
        product_norms[order_id] = []
        
        for product_id in order_data['STOCK_EX_LENH_SAN_XUAT_CHI_TIET_THANH_PHAM_IDS']:
            # Lấy chi tiết sản phẩm
            product_params = {
                "model": "stock.ex.lenh.san.xuat.chi.tiet.thanh.pham",
                "method": "read",
                "args": [product_id, ['MA_HANG_ID', 'TEN_THANH_PHAM', 'SO_LUONG', 'STOCK_EX_THANH_PHAM_CHI_TIET_DINH_MUC_XUAT_NVL_IDS']],
                "kwargs": {'active_menu': 171},
            }
            
            product_details, session = fetch_erp_data(endpoint='call_kw/stock.ex.lenh.san.xuat.chi.tiet.thanh.pham/read', params=product_params, session=session) # Gọi API lấy chi tiết sản phẩm
            # Retry nếu lần đầu fail
            if not product_details: 
                product_details, session = fetch_erp_data(endpoint='call_kw/stock.ex.lenh.san.xuat.chi.tiet.thanh.pham/read', params=product_params)
            
            product_details = product_details[0]
            
            # Lấy định mức nguyên vật liệu
            material_details = []
            if product_details['STOCK_EX_THANH_PHAM_CHI_TIET_DINH_MUC_XUAT_NVL_IDS']:
                material_params = {
                    "model": "stock.ex.thanh.pham.chi.tiet.dinh.muc.xuat.nvl",
                    "method": 'read',
                    "args": [product_details['STOCK_EX_THANH_PHAM_CHI_TIET_DINH_MUC_XUAT_NVL_IDS'], ['MA_HANG_ID', 'SO_LUONG_NVL']],
                    "kwargs": {'active_menu': 171},
                }
                # Gọi API lấy chi tiết sản phẩm
                material_details, session = fetch_erp_data(endpoint='call_kw/stock.ex.thanh.pham.chi.tiet.dinh.muc.xuat.nvl/read', params=material_params, session=session) 

            # Tổ chức dữ liệu
            product_norms[order_id].append({
                'id_san_pham': product_details['MA_HANG_ID'][1],
                'ten_san_pham': product_details['TEN_THANH_PHAM'],
                'so_luong_san_pham': product_details['SO_LUONG'],
                'dinh_muc_nvl': {
                    item['id']: {
                        "id_nguyen_vat_lieu": item['MA_HANG_ID'][1],
                        "dinh_muc_nvl": item['SO_LUONG_NVL']
                    }
                    for item in material_details or []
                }
            })
            
            all_detail_ids.extend([item['id'] for item in material_details or []])

    return production_orders, product_norms, all_detail_ids


def sync_database(production_orders, product_norms, all_detail_ids):
    """Đồng bộ dữ liệu với database: thêm/sửa/xóa"""
    # 1. Lấy dữ liệu hiện tại từ database
    existing_orders = LenhSanXuat.objects.all()
    existing_details = CtLenhSanXuat.objects.all()
    
    existing_order_ids = set(existing_orders.values_list('id_lenh_san_xuat', flat=True))
    existing_detail_ids = set(existing_details.values_list('id_ct_lenh_san_xuat', flat=True))
    
    # 2. So sánh và phân loại IDs (mới/cũ/xóa)
    new_order_ids = set(production_orders.keys()) - existing_order_ids  # Lệnh mới từ Cloudify
    update_order_ids = existing_order_ids & set(production_orders.keys())  # Lệnh cần cập nhật
    new_detail_ids = set(all_detail_ids) - existing_detail_ids  # Chi tiết mới
    update_detail_ids = existing_detail_ids & set(all_detail_ids)  # Chi tiết cần cập nhật
    delete_detail_ids = existing_detail_ids - set(all_detail_ids)  # Chi tiết cần xóa

    # 3. Chuẩn bị containers cho objects
    new_orders, new_details, update_orders, update_details = [], [], [], []
    
    # Tạo dict để tra cứu nhanh
    existing_orders_dict = {o.id_lenh_san_xuat: o for o in existing_orders}
    existing_details_dict = {d.id_ct_lenh_san_xuat: d for d in existing_details}
    order_objects = {}  # Cache các order object mới tạo

    # 4. Tạo các lệnh sản xuất mới
    for order_id in new_order_ids:
        order_data = production_orders[order_id]
        order_obj = LenhSanXuat(
            id_lenh_san_xuat=order_id,
            id_don_hang=order_data['LAP_TU_DON_DAT_HANG_IDS'][0] if order_data['LAP_TU_DON_DAT_HANG_IDS'] else None,
            ngay_tao_don_hang=order_data['NGAY'].split(" ")[0]
        )
        new_orders.append(order_obj)
        order_objects[order_id] = order_obj  # Cache để dùng cho details

    # 5. Cập nhật các lệnh sản xuất cũ
    for order_id in update_order_ids:
        order_obj = existing_orders_dict[order_id]
        order_data = production_orders[order_id]
        order_obj.id_don_hang = order_data['LAP_TU_DON_DAT_HANG_IDS'][0] if order_data['LAP_TU_DON_DAT_HANG_IDS'] else None
        update_orders.append(order_obj)

    # 6. Xử lý chi tiết lệnh sản xuất
    for order_id, products in product_norms.items():
        # Lấy reference đến order (mới hoặc cũ)
        order_ref = order_objects.get(order_id) or existing_orders_dict.get(order_id)
        
        for product in products:
            for detail_id, detail_data in product['dinh_muc_nvl'].items():
                # Chuẩn bị data cho detail object
                detail_obj_data = {
                    'id_ct_lenh_san_xuat': detail_id,
                    'id_lenh_san_xuat': order_ref,
                    'id_san_pham_id': product['id_san_pham'],
                    'ten_san_pham': product['ten_san_pham'],
                    'so_luong_san_pham': product['so_luong_san_pham'],
                    'id_nguyen_vat_lieu_id': detail_data['id_nguyen_vat_lieu'],
                    'so_luong_nguyen_vat_lieu': detail_data['dinh_muc_nvl'],
                }
                
                # Phân loại: tạo mới hoặc cập nhật
                if detail_id in new_detail_ids:
                    new_details.append(CtLenhSanXuat(**detail_obj_data))
                elif detail_id in update_detail_ids:
                    detail_obj = existing_details_dict[detail_id]
                    for key, value in detail_obj_data.items():
                        setattr(detail_obj, key, value)
                    update_details.append(detail_obj)

    # 7. Thực hiện database operations trong transaction
    with transaction.atomic():
        # Tạo mới
        LenhSanXuat.objects.bulk_create(new_orders)
        CtLenhSanXuat.objects.bulk_create(new_details)
        
        # Cập nhật
        LenhSanXuat.objects.bulk_update(update_orders, ['id_don_hang'])
        CtLenhSanXuat.objects.bulk_update(update_details, [
            'id_lenh_san_xuat', 'id_san_pham_id', 'ten_san_pham', 'so_luong_san_pham', 
            'id_nguyen_vat_lieu_id', 'so_luong_nguyen_vat_lieu'
        ])
        
        # Xóa các chi tiết không còn cần thiết
        existing_details.filter(id_ct_lenh_san_xuat__in=delete_detail_ids).delete()

    return len(new_orders), len(update_orders), len(delete_detail_ids)


@require_POST
def orders_sync_cloudify(request):
    """Xử lý đồng bộ hóa dữ liệu đơn hàng từ Cloudify"""
    try:
        # Lấy dữ liệu từ Cloudify
        production_orders, product_norms, all_detail_ids = fetch_cloudify_orders_data()
        
        if not production_orders:
            return JsonResponse({
                'success': False,
                'message': 'Không có dữ liệu lệnh sản xuất từ Cloudify!'
            }, status=404)

        # Đồng bộ với database
        created, updated, deleted = sync_database(production_orders, product_norms, all_detail_ids)

        return JsonResponse({
            'success': True,
            'message': f'Đồng bộ thành công! Tạo: {created}, Sửa: {updated}, Xóa: {deleted}',
            'data_lenh_san_xuat': production_orders,
            'data_product_norms': product_norms
        })

    except Exception as e:
        return JsonResponse({
            'success': False,
            'message': f'Lỗi đồng bộ hóa: {str(e)}'
        }, status=500)


def orders_export(request, pk):
    """Xuất báo cáo tỉ lệ phối trộn theo định dạng PDF hoặc Excel."""
    export_format = request.GET.get('format', '').lower()
    
    # Lấy và tổ chức dữ liệu
    try:
        order_data = get_order_data(pk)
    except Exception as e:
        return JsonResponse({'success': False, 'message': f'Dữ liệu thiếu hoặc không hợp lệ: {e}'}, status=400)
    
    context = {**order_data, 'today': timezone.now()}
    
    if export_format == 'pdf':
        return render(request, 'form/ti_le_dau_tron_pdf.html', context)
    elif export_format == 'excel':
        return create_ti_le_dau_tron_excel_response(pk, context)
    else:
        return JsonResponse({'success': False, 'message': 'Format không hợp lệ!'}, status=400)

def get_order_data(order_id):
    """Lấy và tổ chức dữ liệu lệnh sản xuất."""
    
    # Lấy chi tiết lệnh sản xuất
    order_items = CtLenhSanXuat.objects.filter(
        id_lenh_san_xuat=order_id
    ).filter(
        ~Q(id_nguyen_vat_lieu__nhom_vthh="NVL - THÔ")
    ).select_related('id_san_pham', 'id_nguyen_vat_lieu')

    
    # Xác định danh sách nguyên liệu để làm header
    material_types = sorted({obj.id_nguyen_vat_lieu.ten_khac if obj.id_nguyen_vat_lieu.ten_khac else obj.id_nguyen_vat_lieu.ten_sp_chinh for obj in order_items })

    # Gom nhóm theo sản phẩm
    products_map = {}
    total_quantity = 0
    material_totals = defaultdict(float)
    
    for item in order_items:
        product_id = item.id_san_pham.id_san_pham
        material_name = item.id_nguyen_vat_lieu.ten_khac or item.id_nguyen_vat_lieu.ten_sp_chinh
        material_qty = item.so_luong_nguyen_vat_lieu or 0
        
        # Khởi tạo thông tin sản phẩm nếu chưa có
        if product_id not in products_map:
            product_qty = item.so_luong_san_pham or 0
            products_map[product_id] = {
                'ten_san_pham': item.ten_san_pham or item.id_san_pham.ten_khac,
                'id_san_pham': item.id_san_pham,
                'so_luong_san_pham': product_qty,
                'materials': defaultdict(float),
                'total_materials': 0.0,
            }
            total_quantity += product_qty
        
        # Cộng dồn NVL
        products_map[product_id]['materials'][material_name] += material_qty
        products_map[product_id]['total_materials'] += material_qty
        material_totals[material_name] += material_qty
    
    # Chuyển defaultdict sang dict thông thường
    for product in products_map.values():
        product['materials'] = dict(product['materials'])
    
    return {
        'order_items': list(products_map.values()),
        'material_types': material_types,
        'totals': dict(material_totals),
        'total_quantity_sanpham': total_quantity,
        'total_quantity_nguyenlieu': sum(material_totals.values()),
    }

def create_ti_le_dau_tron_excel_response(order_id, data):
    """Tạo file Excel báo cáo tỉ lệ phối trộn."""
    wb = Workbook()
    ws = wb.active
    
    # Thiết lập styles cho Excel
    styles = {
        'border': Border(left=Side(style='thin'), right=Side(style='thin'), 
                        top=Side(style='thin'), bottom=Side(style='thin')),
        'fill': PatternFill(start_color="F0F0F0", end_color="F0F0F0", fill_type="solid"),
        'base_font': Font(name='Times New Roman', size=12),
        'bold_font': Font(name='Times New Roman', size=12, bold=True),
        'title_font': Font(name='Times New Roman', size=14, bold=True),
        'center': Alignment(horizontal='center', vertical='center'),
        'right': Alignment(horizontal='right')
    }
    
    # ===== PHẦN 1: THÔNG TIN CÔNG TY =====
    company_info = [
        "CÔNG TY CỔ PHẦN TÂN PHONG", 
        "Địa chỉ: Thị Trấn Hùng Sơn, Huyện Lâm Thao, Tỉnh Phú Thọ", 
        "Điện thoại: 0210 221 5277", 
        "Mã số thuế: 2600274542"
    ]
    for i, info in enumerate(company_info):
        cell = ws.cell(row=i+1, column=1, value=info)
        cell.font = styles['bold_font'] if i == 0 else styles['base_font']
    
    # ===== PHẦN 2: TIÊU ĐỀ VÀ HEADER =====
    material_types = data['material_types']
    total_cols = len(material_types) + 5  # Tổng số cột
    
    # Tiêu đề bảng
    ws.merge_cells(start_row=6, start_column=1, end_row=6, end_column=total_cols)
    title_cell = ws.cell(row=6, column=1, value="BẢNG TỈ LỆ PHỐI TRỘN CÁC MẶT HÀNG CHÈ")
    apply_cell_style(title_cell, font=styles['title_font'], align=styles['center'])

    # Header dòng 1
    headers = ["STT", "Tên Hàng Hoá", "Mã HS", "Số lượng", "Thành phần"] + material_types[1:] + ["Tổng cộng"]
    for i, header in enumerate(headers, 1):
        if header:
            cell = ws.cell(row=7, column=i, value=header)
            apply_cell_style(cell, font=styles['bold_font'], border=styles['border'], 
                         align=styles['center'], fill=styles['fill'])

    # Merge cột thành phần
    ws.merge_cells(start_row=7, start_column=5, end_row=7, end_column=5 + len(material_types) - 1)

    # Header dòng 2 (thành phần chi tiết)
    for i, material in enumerate(material_types):
        cell = ws.cell(row=8, column=i+5, value=material)
        apply_cell_style(cell, font=styles['bold_font'], border=styles['border'], 
                     align=styles['center'], fill=styles['fill'])

    # Merge các cột cơ bản
    for col in [1, 2, 3, 4, len(headers)]:
        ws.merge_cells(start_row=7, start_column=col, end_row=8, end_column=col)
    
    # ===== PHẦN 3: DỮ LIỆU SẢN PHẨM =====
    material_cols = {m: i+5 for i, m in enumerate(material_types)}
    start_row = 9
    order_items = data['order_items']
    
    for row_idx, item in enumerate(order_items, start_row):
        # STT, Tên HH, Mã HS, Số lượng
        basic_cells = [
            (1, row_idx-start_row+1),
            (2, item['ten_san_pham']),
            (3, item['id_san_pham'].ma_hs),
            (4, item['so_luong_san_pham'])
        ]
        
        for col, value in basic_cells:
            cell = ws.cell(row=row_idx, column=col, value=value)
            number_format = '#,##0.00' if col == 4 else None
            apply_cell_style(cell, font=styles['base_font'], border=styles['border'], 
                         align=styles['center'], number_format=number_format)
        
        # Nguyên liệu
        for mat, col in material_cols.items():
            val = item['materials'].get(mat)
            cell = ws.cell(row=row_idx, column=col, value=val)
            apply_cell_style(cell, font=styles['base_font'], border=styles['border'], 
                         align=styles['center'], number_format='#,##0.00' if val is not None else None)
        
        # Tổng cộng
        cell = ws.cell(row=row_idx, column=total_cols, value=item['total_materials'])
        apply_cell_style(cell, font=styles['base_font'], border=styles['border'], 
                     align=styles['center'], number_format='#,##0.00')
    
    # ===== PHẦN 4: DÒNG TỔNG CỘNG =====
    total_row = start_row + len(order_items)
    
    # Tiêu đề "Tổng cộng"
    cell = ws.cell(row=total_row, column=1, value="Tổng cộng")
    apply_cell_style(cell, font=styles['bold_font'], border=styles['border'], align=styles['center'])
    ws.merge_cells(start_row=total_row, start_column=1, end_row=total_row, end_column=3)

    # Tổng số lượng sản phẩm
    cell = ws.cell(row=total_row, column=4, value=data['total_quantity_sanpham'])
    apply_cell_style(cell, font=styles['bold_font'], border=styles['border'], 
                 align=styles['center'], number_format='#,##0.00')

    # Tổng lượng nguyên liệu
    for key, col in material_cols.items():
        cell = ws.cell(row=total_row, column=col, value=data['totals'].get(key))
        apply_cell_style(cell, font=styles['bold_font'], border=styles['border'], 
                     align=styles['center'], number_format='#,##0.00')
    
    # Tổng cộng cuối cùng
    cell = ws.cell(row=total_row, column=total_cols, value=data['total_quantity_nguyenlieu'])
    apply_cell_style(cell, font=styles['bold_font'], border=styles['border'], 
                 align=styles['center'], number_format='#,##0.00')
    
    # ===== PHẦN 5: CHỮ KÝ VÀ CAM KẾT =====
    sig_row = total_row + 3
    today = data['today']

    # Cam kết
    commitment = "Công ty cam kết số liệu, thông tin khai báo trên là đúng và chịu trách nhiệm trước pháp luật về thông tin, số liệu đã khai."
    cell = ws.cell(row=sig_row, column=1, value=commitment)
    cell.font = styles['base_font']
    ws.merge_cells(start_row=sig_row, start_column=1, end_row=sig_row, end_column=total_cols)

    # Ngày tháng, người đại diện, chỗ ký
    signatures = [
        (sig_row+2, total_cols-1, f"Ngày {today.day} tháng {today.month} năm {today.year}", 
         styles['base_font']),
        (sig_row+4, total_cols, "NGƯỜI ĐẠI DIỆN THEO PHÁP LUẬT CỦA THƯƠNG NHÂN", 
         styles['bold_font']),
        (sig_row+6, total_cols-1, "(Ký, đóng dấu, ghi rõ họ, tên)", 
         styles['base_font'])
    ]

    for row, col, text, font in signatures:
        cell = ws.cell(row=row, column=col, value=text)
        apply_cell_style(cell, font=font, align=styles['right'])
    
    # ===== PHẦN 6: ĐIỀU CHỈNH CHIỀU RỘNG CỘT =====
    # Chỉ xem xét các dòng quan trọng để tối ưu hiệu suất
    for column in ws.columns:
        max_length = 0
        column_letter = get_column_letter(column[0].column)
        
        for cell in column[6:6+len(order_items)]:  # Chỉ kiểm tra các dòng header và một vài dòng đầu tiên
            try:
                cell_length = len(str(cell.value)) if cell.value else 0
                if cell_length > max_length:
                    max_length = cell_length
            except:
                pass
        
        ws.column_dimensions[column_letter].width = max_length + 3
    
    # ===== TẠO RESPONSE =====
    output = BytesIO()
    wb.save(output)
    output.seek(0)
    
    filename = f"Lenh_SX_{order_id}_{today.strftime('%Y-%m-%d')}.xlsx"
    response = HttpResponse(
        output.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response

def apply_cell_style(cell, font=None, border=None, align=None, fill=None, number_format=None):
    """Áp dụng style cho một ô Excel."""
    if font: cell.font = font
    if border: cell.border = border
    if align: cell.alignment = align
    if fill: cell.fill = fill
    if number_format and cell.value is not None: cell.number_format = number_format

# ==================== HELPER FUNCTIONS ====================
def _get_vat_tu_info(id_san_pham):
    """
    Get VatTu information by product ID - only for NVL-KHÔ items
    Returns tuple: (vat_tu_object, ty_le_thu_hoi, don_vi_tinh, ten_nguyen_lieu, ma_hs)
    """
    try:
        if id_san_pham:
            vat_tu = VatTu.objects.get(id_san_pham=id_san_pham, nhom_vthh='NVL - KHÔ')
            return (
                vat_tu,
                vat_tu.ty_le_thu_hoi,
                vat_tu.don_vi_tinh or 'kg',
                vat_tu.ten_khac,
                vat_tu.ma_hs
            )
    except VatTu.DoesNotExist:
        pass
    return None, None, 'kg', None, ''


def _calculate_quantities(record, ty_le_thu_hoi):
    """
    Calculate derived quantities for rollback record
    Returns tuple: (so_luong_thanh_pham_thu_hoi, so_luong_thanh_pham_ton_kho)
    """
    # Calculate finished product recovery quantity
    so_luong_thanh_pham_thu_hoi = record.thanh_pham_thu_hoi
    if (so_luong_thanh_pham_thu_hoi is None and 
        record.so_luong_san_xuat is not None and 
        ty_le_thu_hoi is not None):
        so_luong_thanh_pham_thu_hoi = record.so_luong_san_xuat * ty_le_thu_hoi
    
    # Calculate inventory quantity
    so_luong_thanh_pham_ton_kho = record.so_luong_thanh_pham_ton_kho
    if (so_luong_thanh_pham_ton_kho is None and 
        so_luong_thanh_pham_thu_hoi is not None and 
        record.so_luong_san_pham_xuat is not None):
        so_luong_thanh_pham_ton_kho = so_luong_thanh_pham_thu_hoi - record.so_luong_san_pham_xuat
    
    return so_luong_thanh_pham_thu_hoi, so_luong_thanh_pham_ton_kho


def _get_production_order_output_quantity(record):
    """Get production order output quantity from CtLenhSanXuat - only for NVL-KHÔ items"""
    so_luong_san_pham_xuat = record.so_luong_san_pham_xuat
    try:
        ct_lenh_sx = CtLenhSanXuat.objects.filter(
            id_lenh_san_xuat=record.id_lenh_san_xuat,
            id_nguyen_vat_lieu__id_san_pham=record.id_san_pham,
            id_nguyen_vat_lieu__nhom_vthh='NVL - KHÔ'
        ).first()
        
        if ct_lenh_sx:
            so_luong_san_pham_xuat = ct_lenh_sx.so_luong_nguyen_vat_lieu
    except CtLenhSanXuat.DoesNotExist:
        pass
    
    return so_luong_san_pham_xuat


def _determine_status(record, so_luong_san_pham_xuat, so_luong_thanh_pham_thu_hoi):
    """Determine the status of a rollback record"""
    if record.trang_thai is not None:
        return record.trang_thai
    
    # Auto-determine status based on data completeness
    if (record.so_luong_mua_vao is not None and 
        record.so_luong_san_xuat is not None and 
        so_luong_san_pham_xuat is not None and 
        so_luong_thanh_pham_thu_hoi is not None):
        return "Hoàn thành"
    else:
        return "Đang xử lý"


def _format_rollback_record(record):
    """
    Format a single rollback record for display
    Returns formatted dictionary
    """
    # Get VatTu information - only for NVL-KHÔ items
    vat_tu, ty_le_thu_hoi, don_vi_tinh, ten_khac, ma_hs = _get_vat_tu_info(record.id_san_pham)
    
    # Use alternative name from VatTu if available
    ten_nguyen_lieu = ten_khac or record.ten_nguyen_lieu
    ma_hs = ma_hs if ma_hs else ''
    
    # Get production order output quantity
    so_luong_san_pham_xuat = _get_production_order_output_quantity(record)
    
    # Calculate derived quantities
    so_luong_thanh_pham_thu_hoi, so_luong_thanh_pham_ton_kho = _calculate_quantities(record, ty_le_thu_hoi)
    
    # Determine status
    trang_thai = _determine_status(record, so_luong_san_pham_xuat, so_luong_thanh_pham_thu_hoi)
    
    return {
        'id_bang_ke_tru_lui': record.id_bang_ke_tru_lui,
        'ma_lenh_sx': record.id_lenh_san_xuat.id_lenh_san_xuat,
        'ma_don_hang': record.id_lenh_san_xuat.id_don_hang,
        'ten_nguyen_lieu': ten_nguyen_lieu,
        'ma_hs': ma_hs,
        'don_vi_tinh': don_vi_tinh,
        'so_luong_mua_vao': record.so_luong_mua_vao,
        'so_luong_san_xuat': record.so_luong_san_xuat,
        'ty_le_thu_hoi': ty_le_thu_hoi,
        'so_luong_thanh_pham_thu_hoi': so_luong_thanh_pham_thu_hoi,
        'so_luong_san_pham_xuat': so_luong_san_pham_xuat,
        'so_luong_thanh_pham_ton_kho': so_luong_thanh_pham_ton_kho,
        'ghi_chu': record.ghi_chu,
        'ngay_thang': record.ngay_thang,
        'trang_thai': trang_thai
    }


def _safe_float_conversion(value, default=0):
    """Safely convert value to float with default fallback"""
    try:
        return float(value) if value not in (None, '', 'null') else default
    except (ValueError, TypeError):
        return default


def _parse_date(date_string):
    """Parse date string to date object (support both YYYY-MM-DD and DD/MM/YYYY)"""
    if not date_string:
        return date.today()
    for fmt in ('%Y-%m-%d', '%d/%m/%Y'):
        try:
            return datetime.strptime(date_string, fmt).date()
        except ValueError:
            continue
    return date.today()


def _update_vat_tu_recovery_rate(id_san_pham, ty_le_thu_hoi):
    """Update recovery rate in VatTu table - only for NVL-KHÔ items"""
    if not id_san_pham or ty_le_thu_hoi <= 0:
        return
    
    try:
        vat_tu = VatTu.objects.get(id_san_pham=id_san_pham, nhom_vthh='NVL - KHÔ')
        vat_tu.ty_le_thu_hoi = ty_le_thu_hoi
        vat_tu.save()
    except VatTu.DoesNotExist:
        pass


def _update_production_order_detail(lenh_sx, id_san_pham, so_luong_san_pham_xuat):
    """Update production order detail quantity - only for NVL-KHÔ items"""
    if not id_san_pham:
        return
    
    try:
        ct_lenh_sx = CtLenhSanXuat.objects.filter(
            id_lenh_san_xuat=lenh_sx,
            id_nguyen_vat_lieu__id_san_pham=id_san_pham,
            id_nguyen_vat_lieu__nhom_vthh='NVL - KHÔ'
        ).first()
        
        if ct_lenh_sx:
            ct_lenh_sx.so_luong_nguyen_vat_lieu = so_luong_san_pham_xuat
            ct_lenh_sx.save()
    except Exception as e:
        print(f"Error updating production order detail: {e}")


# ==================== ROLLBACK MAIN VIEWS ====================
def rollback(request):
    """Main rollback list view with filtering - only show NVL-KHÔ related records"""
    # Get filter parameters
    ma_don_hang = request.GET.get('ma_don_hang')
    ma_lenh_sx = request.GET.get('ma_lenh_sx')
    
    records = BangKeTruLuiNguyenLieu.objects.select_related('id_lenh_san_xuat').all()
    
    # Apply filters
    if ma_don_hang:
        records = records.filter(id_lenh_san_xuat__id_don_hang=ma_don_hang)
    
    if ma_lenh_sx:
        records = records.filter(id_lenh_san_xuat__id_lenh_san_xuat=ma_lenh_sx)
    
    # Format records for display
    formatted_records = [_format_rollback_record(record) for record in records]
    
    # Get order list for dropdown
    don_hang_list = LenhSanXuat.objects.values_list('id_don_hang', flat=True).distinct()
    
    context = {
        'records': formatted_records,
        'don_hang_list': don_hang_list
    }
    
    return render(request, 'rollback.html', context)


def rollback_detail(request, pk):
    """Rollback detail view - only for NVL-KHÔ items"""
    try:
        # Filter for NVL-KHÔ items only
        record = BangKeTruLuiNguyenLieu.objects.select_related('id_lenh_san_xuat').get(pk=pk)
        formatted_record = _format_rollback_record(record)
        context = {'record': formatted_record}
    except BangKeTruLuiNguyenLieu.DoesNotExist:
        context = {'record': None}
    
    return render(request, 'rollback_detail.html', context)


def rollback_create(request):
    """Create new rollback records - only for NVL-KHÔ items"""
    don_hang_list = LenhSanXuat.objects.values_list('id_don_hang', flat=True).distinct()
    
    if request.method == 'GET':
        return render(request, 'rollback_create.html', {'don_hang_list': don_hang_list})
    
    # POST request processing
    ma_lenh_sx = request.POST.get('ma_lenh_sx')
    
    if not ma_lenh_sx:
        return render(request, 'rollback_create.html', {
            'error': 'Vui lòng chọn lệnh sản xuất',
            'don_hang_list': don_hang_list
        })
    
    # Get production order
    try:
        lenh_sx = LenhSanXuat.objects.get(id_lenh_san_xuat=ma_lenh_sx)
    except LenhSanXuat.DoesNotExist:
        return render(request, 'rollback_create.html', {
            'error': f'Không tìm thấy lệnh sản xuất với mã: {ma_lenh_sx}',
            'don_hang_list': don_hang_list
        })
    
    # Get form data lists
    form_fields = [
        'id_san_pham', 'ten_nguyen_lieu', 'ma_hs', 'don_vi_tinh',
        'so_luong_mua_vao', 'so_luong_san_xuat', 'ty_le_thu_hoi',
        'sl_thanh_pham_thu_hoi', 'so_luong_san_pham_xuat',
        'sl_thanh_pham_ton_kho', 'ghi_chu', 'ngay_tao'
    ]
    
    form_data = {}
    for field in form_fields:
        form_data[field] = request.POST.getlist(f'{field}[]')
    
    if not form_data['ten_nguyen_lieu']:
        return render(request, 'rollback_create.html', {
            'error': 'Không có dữ liệu nguyên liệu để lưu',
            'don_hang_list': don_hang_list
        })
    
    # Process each material - only NVL-KHÔ items
    created_records = []
    for i in range(len(form_data['ten_nguyen_lieu'])):
        try:
            # Check if the product is NVL-KHÔ before processing
            id_san_pham = form_data['id_san_pham'][i] if i < len(form_data['id_san_pham']) else None
            if id_san_pham:
                try:
                    VatTu.objects.get(id_san_pham=id_san_pham, nhom_vthh='NVL - KHÔ')
                except VatTu.DoesNotExist:
                    # Skip this record if it's not NVL-KHÔ
                    continue
            
            record_data = _process_single_material_record(form_data, i, lenh_sx)
            if record_data:
                bangke_tru_lui = BangKeTruLuiNguyenLieu(**record_data)
                bangke_tru_lui.save()
                created_records.append(bangke_tru_lui)
                
                # Update related data
                ty_le_thu_hoi = _safe_float_conversion(
                    form_data['ty_le_thu_hoi'][i] if i < len(form_data['ty_le_thu_hoi']) else None
                )
                if ty_le_thu_hoi > 1:
                    ty_le_thu_hoi = ty_le_thu_hoi / 100
                
                _update_vat_tu_recovery_rate(id_san_pham, ty_le_thu_hoi)
                _update_production_order_detail(
                    lenh_sx, 
                    id_san_pham, 
                    record_data.get('so_luong_san_pham_xuat', 0)
                )
                
        except Exception as e:
            print(f"Error processing material record {i}: {e}")
            continue
    
    if not created_records:
        return render(request, 'rollback_create.html', {
            'error': 'Không thể tạo bảng kê trừ lùi. Vui lòng kiểm tra lại dữ liệu hoặc đảm bảo có ít nhất một nguyên liệu thuộc nhóm NVL - KHÔ.',
            'don_hang_list': don_hang_list
        })
    
    return redirect('rollback')


def _process_single_material_record(form_data, index, lenh_sx):
    """Process a single material record from form data - only for NVL-KHÔ items"""
    i = index
    
    # Get basic data
    id_san_pham = form_data['id_san_pham'][i] if i < len(form_data['id_san_pham']) else None
    ten_nguyen_lieu = form_data['ten_nguyen_lieu'][i]
    ma_hs = form_data['ma_hs'][i] if i < len(form_data['ma_hs']) else ''
    
    # Get ma_hs from VatTu if not provided - only for NVL-KHÔ items
    if not ma_hs and id_san_pham:
        _, _, _, _, ma_hs = _get_vat_tu_info(id_san_pham)
    
    # Parse quantities
    so_luong_mua_vao = _safe_float_conversion(
        form_data['so_luong_mua_vao'][i] if i < len(form_data['so_luong_mua_vao']) else None
    )
    so_luong_san_xuat = _safe_float_conversion(
        form_data['so_luong_san_xuat'][i] if i < len(form_data['so_luong_san_xuat']) else None
    )
    
    # Parse recovery rate
    ty_le_thu_hoi = _safe_float_conversion(
        form_data['ty_le_thu_hoi'][i] if i < len(form_data['ty_le_thu_hoi']) else None
    )
    if ty_le_thu_hoi > 1:
        ty_le_thu_hoi = ty_le_thu_hoi / 100
    
    # Calculate or get recovery quantity
    if i < len(form_data['sl_thanh_pham_thu_hoi']) and form_data['sl_thanh_pham_thu_hoi'][i]:
        so_luong_thanh_pham_thu_hoi = _safe_float_conversion(form_data['sl_thanh_pham_thu_hoi'][i])
    else:
        so_luong_thanh_pham_thu_hoi = so_luong_san_xuat * ty_le_thu_hoi
    
    # Get output quantity
    so_luong_san_pham_xuat = _safe_float_conversion(
        form_data['so_luong_san_pham_xuat'][i] if i < len(form_data['so_luong_san_pham_xuat']) else None
    )
    
    # Calculate or get inventory quantity
    if i < len(form_data['sl_thanh_pham_ton_kho']) and form_data['sl_thanh_pham_ton_kho'][i]:
        so_luong_ton_kho = _safe_float_conversion(form_data['sl_thanh_pham_ton_kho'][i])
    else:
        so_luong_ton_kho = so_luong_thanh_pham_thu_hoi - so_luong_san_pham_xuat
    
    # Parse date
    ngay_tao = _parse_date(
        form_data['ngay_tao'][i] if i < len(form_data['ngay_tao']) and form_data['ngay_tao'][i] else None
    )
    
    # Determine status
    trang_thai = "Hoàn thành" if so_luong_mua_vao > 0 and so_luong_san_xuat > 0 else "Đang xử lý"
    
    # Get notes
    ghi_chu = form_data['ghi_chu'][i] if i < len(form_data['ghi_chu']) else ''
    
    return {
        'id_lenh_san_xuat': lenh_sx,
        'id_san_pham': id_san_pham,
        'ten_nguyen_lieu': ten_nguyen_lieu,
        'so_luong_mua_vao': so_luong_mua_vao,
        'so_luong_san_xuat': so_luong_san_xuat,
        'thanh_pham_thu_hoi': so_luong_thanh_pham_thu_hoi,
        'so_luong_san_pham_xuat': so_luong_san_pham_xuat,
        'so_luong_thanh_pham_ton_kho': so_luong_ton_kho,
        'ngay_thang': ngay_tao,
        'trang_thai': trang_thai,
        'ghi_chu': ghi_chu
    }

# ==================== ROLLBACK UPDATE VIEW ====================
@csrf_exempt
@require_POST
def rollback_update(request, pk):
    """
    API để cập nhật số lượng mua vào và sản xuất,
    các giá trị tự tính như thu hồi, tồn kho, trạng thái được cập nhật tự động.
    """
    try:
        record = BangKeTruLuiNguyenLieu.objects.get(pk=pk)
        data = json.loads(request.body)

        # Nhập số lượng mới từ người dùng
        so_luong_mua_vao = float(data.get('so_luong_mua_vao', 0))
        so_luong_san_xuat = float(data.get('so_luong_san_xuat', 0))

        # Lấy tỷ lệ thu hồi từ bảng nguyên vật liệu
        try:
            vat_tu = VatTu.objects.get(id_san_pham=record.id_san_pham, nhom_vthh='NVL - KHÔ')
            ty_le_thu_hoi = vat_tu.ty_le_thu_hoi or 0
        except VatTu.DoesNotExist:
            ty_le_thu_hoi = 0

        # Tính lại dữ liệu
        thanh_pham_thu_hoi = so_luong_san_xuat * ty_le_thu_hoi
        so_luong_san_pham_xuat = record.so_luong_san_pham_xuat or 0
        so_luong_thanh_pham_ton_kho = thanh_pham_thu_hoi - so_luong_san_pham_xuat

        # Xác định trạng thái
        if so_luong_mua_vao > 0 and so_luong_san_xuat > 0 and thanh_pham_thu_hoi >= 0:
            trang_thai = "Hoàn thành"
        else:
            trang_thai = "Đang xử lý"

        # Cập nhật bản ghi
        record.so_luong_mua_vao = so_luong_mua_vao
        record.so_luong_san_xuat = so_luong_san_xuat
        record.thanh_pham_thu_hoi = thanh_pham_thu_hoi
        record.so_luong_thanh_pham_ton_kho = so_luong_thanh_pham_ton_kho
        record.trang_thai = trang_thai

        record.save()

        return JsonResponse({'status': 'success'})
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)})

# ==================== ROLLBACK DELETE VIEW ====================
@csrf_exempt
@require_http_methods(["DELETE"])
def rollback_delete(request, pk):
    """
    API để xóa bảng kê trừ lùi nguyên liệu
    """
    try:
        record = BangKeTruLuiNguyenLieu.objects.get(pk=pk)
        # Lưu thông tin để log
        ten_nguyen_lieu = record.ten_nguyen_lieu
        ma_lenh_sx = record.id_lenh_san_xuat.id_lenh_san_xuat
        
        # Xóa record
        record.delete()
        
        print(f"Đã xóa bảng kê trừ lùi: {ten_nguyen_lieu} - Lệnh SX: {ma_lenh_sx}")
        
        return JsonResponse({
            'status': 'success',
            'message': f'Đã xóa bảng kê "{ten_nguyen_lieu}" thành công'
        })
        
    except BangKeTruLuiNguyenLieu.DoesNotExist:
        return JsonResponse({
            'status': 'error', 
            'message': 'Không tìm thấy bảng kê hoặc sản phẩm không thuộc nhóm NVL - KHÔ'
        })
    except Exception as e:
        print(f"Error deleting rollback record {pk}: {e}")
        return JsonResponse({
            'status': 'error', 
            'message': f'Lỗi khi xóa: {str(e)}'
        })

# ==================== API ENDPOINTS ====================
@require_GET
def get_lenh_san_xuat_all(request):
    """API trả về toàn bộ lệnh sản xuất theo mã đơn hàng (dùng cho filter rollback)"""
    ma_don_hang = request.GET.get('ma_don_hang')
    
    if not ma_don_hang:
        return JsonResponse({'lenh_sx_list': []})

    lenh_sx_list = LenhSanXuat.objects.filter(
        id_don_hang=ma_don_hang
    ).values_list('id_lenh_san_xuat', flat=True)
    
    return JsonResponse({'lenh_sx_list': list(lenh_sx_list)})

@require_GET
def get_lenh_san_xuat(request):
    """API endpoint to get production orders by order code - exclude orders with existing rollback records"""
    ma_don_hang = request.GET.get('ma_don_hang')
    
    if ma_don_hang:
        # Lấy tất cả lệnh sản xuất theo đơn hàng
        all_lenh_sx = LenhSanXuat.objects.filter(
            id_don_hang=ma_don_hang
        ).values_list('id_lenh_san_xuat', flat=True)
        
        # Lấy danh sách lệnh sản xuất đã có bảng kê trừ lùi (chỉ NVL-KHÔ)
        existing_lenh_sx = BangKeTruLuiNguyenLieu.objects.filter(
            id_lenh_san_xuat__id_don_hang=ma_don_hang,
            id_san_pham__in=VatTu.objects.filter(nhom_vthh='NVL - KHÔ').values_list('id_san_pham', flat=True)
        ).values_list('id_lenh_san_xuat__id_lenh_san_xuat', flat=True).distinct()
        
        # Lọc ra các lệnh sản xuất chưa có bảng kê
        available_lenh_sx = [lenh for lenh in all_lenh_sx if lenh not in existing_lenh_sx]
        
        return JsonResponse({
            'lenh_sx_list': list(available_lenh_sx),
            'existing_lenh_sx': list(existing_lenh_sx)  # Trả về để có thể hiển thị thông báo nếu cần
        })
    
    return JsonResponse({'lenh_sx_list': [], 'existing_lenh_sx': []})


@require_GET
def get_lenh_san_xuat_detail(request):
    """API endpoint to get production order details - only for NVL-KHÔ items, gộp theo tên_khac"""
    ma_lenh_sx = request.GET.get('ma_lenh_sx')
    
    if not ma_lenh_sx:
        return JsonResponse({
            'status': 'error',
            'message': 'Vui lòng cung cấp mã lệnh sản xuất'
        })
    
    
    try:
        lenh_sx = LenhSanXuat.objects.get(id_lenh_san_xuat=ma_lenh_sx)

        ct_lenh_sx = CtLenhSanXuat.objects.filter(
            id_lenh_san_xuat=lenh_sx,
            id_nguyen_vat_lieu__nhom_vthh='NVL - KHÔ'
        )

        # Gộp theo tên_khac (tên nguyên liệu)
        nguyen_lieu_dict = {}

        for item in ct_lenh_sx:
            vat_tu = item.id_nguyen_vat_lieu

            if not vat_tu or vat_tu.nhom_vthh != 'NVL - KHÔ':
                continue

            ten_nguyen_lieu = vat_tu.ten_khac or vat_tu.ten_sp_chinh
            if not ten_nguyen_lieu:
                continue

            key = ten_nguyen_lieu.strip().lower()

            if key not in nguyen_lieu_dict:
                nguyen_lieu_dict[key] = {
                    'id_ct_lenh_san_xuat': item.id_ct_lenh_san_xuat,
                    'ten_nguyen_lieu': ten_nguyen_lieu,
                    'ma_hs': vat_tu.ma_hs or '',
                    'don_vi_tinh': vat_tu.don_vi_tinh or 'kg',
                    'ty_le_thu_hoi': vat_tu.ty_le_thu_hoi,
                    'so_luong_nguyen_vat_lieu': item.so_luong_nguyen_vat_lieu or 0,
                    'so_luong_san_pham': item.so_luong_san_pham or 0,
                    'id_san_pham': vat_tu.id_san_pham,
                }
            else:
                # Cộng dồn số lượng
                nguyen_lieu_dict[key]['so_luong_nguyen_vat_lieu'] += item.so_luong_nguyen_vat_lieu or 0
                nguyen_lieu_dict[key]['so_luong_san_pham'] += item.so_luong_san_pham or 0

                # Cập nhật tỷ lệ thu hồi nếu ban đầu chưa có
                if nguyen_lieu_dict[key]['ty_le_thu_hoi'] is None and vat_tu.ty_le_thu_hoi is not None:
                    nguyen_lieu_dict[key]['ty_le_thu_hoi'] = vat_tu.ty_le_thu_hoi

        return JsonResponse({
            'status': 'success',
            'ma_don_hang': lenh_sx.id_don_hang,
            'ma_lenh_sx': lenh_sx.id_lenh_san_xuat,
            'ngay_tao': lenh_sx.ngay_tao_don_hang.strftime('%Y-%m-%d') if lenh_sx.ngay_tao_don_hang else None,
            'chi_tiet': list(nguyen_lieu_dict.values())
        })

    except LenhSanXuat.DoesNotExist:
        return JsonResponse({
            'status': 'error',
            'message': f'Không tìm thấy lệnh sản xuất với mã: {ma_lenh_sx}'
        })
    
@require_GET
def check_rollback_exist(request):
    """API endpoint to check if rollback records exist for a production order - only for NVL-KHÔ items"""
    ma_lenh_sx = request.GET.get('ma_lenh_sx')
    
    if not ma_lenh_sx:
        return JsonResponse({
            'status': 'error',
            'message': 'Vui lòng cung cấp mã lệnh sản xuất'
        })
    
    # Kiểm tra xem lệnh sản xuất đã có bảng kê trừ lùi chưa - chỉ NVL-KHÔ
    existing_rollback = BangKeTruLuiNguyenLieu.objects.filter(
        id_lenh_san_xuat__id_lenh_san_xuat=ma_lenh_sx,
        id_san_pham__in=VatTu.objects.filter(nhom_vthh='NVL - KHÔ').values_list('id_san_pham', flat=True)
    ).exists()
    
    return JsonResponse({
        'exists': existing_rollback,
        'message': f'Lệnh sản xuất {ma_lenh_sx} đã có bảng kê trừ lùi. Vui lòng xóa các bảng kê hiện tại trước khi tạo mới.' if existing_rollback else ''
    })

# ==================== EXPORT FUNCTIONS ====================
@require_GET
def rollback_export_pdf(request, pk):
    """Render rollback record as HTML for PDF printing (AJAX) - only for NVL-KHÔ items"""
    try:
        record = BangKeTruLuiNguyenLieu.objects.filter(
            id_bang_ke_tru_lui=pk,
            id_san_pham__in=VatTu.objects.filter(nhom_vthh='NVL - KHÔ').values_list('id_san_pham', flat=True)
        ).select_related('id_lenh_san_xuat').get()
    except BangKeTruLuiNguyenLieu.DoesNotExist:
        return HttpResponse("Không tìm thấy bảng kê trừ lùi hoặc sản phẩm không thuộc nhóm NVL - KHÔ", status=404)

    formatted_record = _format_rollback_record(record)

    if formatted_record['ty_le_thu_hoi']:
        formatted_record['ty_le_thu_hoi'] *= 100

    ton_kho = int(formatted_record['so_luong_thanh_pham_ton_kho']) if formatted_record['so_luong_thanh_pham_ton_kho'] else 0
    ton_kho_chu = convert_number_to_vietnamese_words(ton_kho)

    context = {
        'record': formatted_record,
        'ton_kho_chu': ton_kho_chu
    }

    return render(request, 'rollback_detail_pdf.html', context)


@require_GET
def rollback_export_excel(request, pk):
    """Export a rollback record to Excel - only for NVL-KHÔ items"""
    try:
        record = BangKeTruLuiNguyenLieu.objects.filter(
            id_bang_ke_tru_lui=pk,
            id_san_pham__in=VatTu.objects.filter(nhom_vthh='NVL - KHÔ').values_list('id_san_pham', flat=True)
        ).select_related('id_lenh_san_xuat').get()
    except BangKeTruLuiNguyenLieu.DoesNotExist:
        return HttpResponse("Không tìm thấy bảng kê trừ lùi hoặc sản phẩm không thuộc nhóm NVL - KHÔ", status=404)
    
    # Generate Excel file
    excel_file = generate_excel(data=None, record=record)
    
    response = HttpResponse(
        excel_file, 
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    response['Content-Disposition'] = f'attachment; filename="bang_ke_tru_lui_{pk}.xlsx"'
    
    return response


# ==================== PURCHASE FROM PEOPLE VIEWS ====================
def get_vat_tu_info_general(id_san_pham):
    """
    Get general VatTu information (for purchase, WO) — không giới hạn nhóm NVL
    Returns: (ten_nguyen_lieu, don_vi_tinh, ma_hs, ty_le_thu_hoi)
    """
    try:
        vat_tu = VatTu.objects.get(id_san_pham=id_san_pham)
        ten_nguyen_lieu = vat_tu.ten_khac or vat_tu.ten_sp_chinh or f"Sản phẩm {id_san_pham}"
        return (
            ten_nguyen_lieu,
            vat_tu.don_vi_tinh or 'kg',
            vat_tu.ma_hs or '',
            vat_tu.ty_le_thu_hoi or 0
        )
    except VatTu.DoesNotExist:
        return (f"Sản phẩm {id_san_pham}", 'kg', '', 0)
    
def purchase(request):
    """Main purchase from people list view with filtering"""
    ma_don_hang = request.GET.get('ma_don_hang')
    ma_lenh_sx = request.GET.get('ma_lenh_sx')
    
    records = BangKeThuMuaTuDan.objects.select_related('id_lenh_san_xuat').all()

    # Apply filters
    if ma_don_hang:
        records = records.filter(id_lenh_san_xuat__id_don_hang=ma_don_hang)
    if ma_lenh_sx:
        records = records.filter(id_lenh_san_xuat__id_lenh_san_xuat=ma_lenh_sx)

    formatted_records = []
    for record in records:
        # Lấy số lượng sản phẩm xuất từ bảng kê trừ lùi
        rollback_record = BangKeTruLuiNguyenLieu.objects.filter(
            id_lenh_san_xuat=record.id_lenh_san_xuat,
            id_san_pham=record.id_san_pham
        ).first()
        so_luong_san_pham_xuat = rollback_record.so_luong_san_pham_xuat if rollback_record else 0

        # Lấy thông tin vật tư và tỷ lệ thu hồi
        ten_nguyen_lieu, _, _, ty_le_thu_hoi = get_vat_tu_info_general(record.id_san_pham)

        # Thời gian thu mua là chuỗi từ model
        thoi_gian_thu_mua = record.ngay_lap_giay_to if record.ngay_lap_giay_to else ''

        formatted_record = {
            'id_bang_ke_thu_mua_tu_dan': record.id_bang_ke_thu_mua_tu_dan,
            'ma_lenh_sx': record.id_lenh_san_xuat.id_lenh_san_xuat,
            'ten_nguyen_lieu': ten_nguyen_lieu,
            'so_luong_san_pham_xuat': so_luong_san_pham_xuat,
            'ty_le_thu_hoi': f"{int(ty_le_thu_hoi * 100)}%" if ty_le_thu_hoi else "0%",
            'so_luong_san_xuat_toi_thieu': int(so_luong_san_pham_xuat / ty_le_thu_hoi) if ty_le_thu_hoi > 0 else 0,
            'thoi_gian_thu_mua': thoi_gian_thu_mua,
            'trang_thai': 'Chi tiết'
        }
        formatted_records.append(formatted_record)

    don_hang_list = LenhSanXuat.objects.values_list('id_don_hang', flat=True).distinct()

    context = {
        'records': formatted_records,
        'don_hang_list': don_hang_list
    }

    return render(request, 'purchase.html', context)


def purchase_detail(request, pk):
    """Purchase from people detail view with edit functionality"""
    try:
        record = BangKeThuMuaTuDan.objects.select_related('id_lenh_san_xuat').get(pk=pk)
        
        if request.method == 'POST':
            # Handle update request
            return handle_purchase_detail_update(request, record)
        
        # GET request - display detail
        # Get detailed records from CtBangKeThuMuaTuDan

        detail_records = CtBangKeThuMuaTuDan.objects.filter(
            id_bang_ke_thu_mua_tu_dan=record
        ).select_related('id_nguoi_ban').order_by('id_ct_bang_ke_thu_mua_tu_dan')
        
        ngay_from = ''
        ngay_to = ''
        if record.ngay_lap_giay_to and ' - ' in record.ngay_lap_giay_to:
            parts = record.ngay_lap_giay_to.split(' - ')
            try:
                ngay_from = datetime.strptime(parts[0], '%d/%m/%Y').date().strftime('%Y-%m-%d')
                ngay_to = datetime.strptime(parts[1], '%d/%m/%Y').date().strftime('%Y-%m-%d')
            except:
                pass

         # Get VatTu info
        ten_nguyen_lieu, don_vi_tinh, _, ty_le_thu_hoi = get_vat_tu_info_general(record.id_san_pham)

        # Format detail records
        formatted_details = []
        for detail in detail_records:
            formatted_detail = {
                'id_ct_bang_ke_thu_mua_tu_dan': detail.id_ct_bang_ke_thu_mua_tu_dan,
                'ten_nguyen_lieu': detail.ten_nguyen_lieu,
                'don_vi_tinh': don_vi_tinh,
                'don_gia': detail.don_gia or 0,
                'so_luong': detail.so_luong or 0,
                'thanh_tien': (detail.don_gia or 0) * (detail.so_luong or 0),
                'ten_nguoi_ban': detail.id_nguoi_ban.ten if detail.id_nguoi_ban else '',
                'id_nguoi_ban': detail.id_nguoi_ban.id if detail.id_nguoi_ban else None,
                'so_cmnd_cccd': detail.id_nguoi_ban.so_cmnd_cccd if detail.id_nguoi_ban else '',
                'ngay_cap_cmnd_cccd': detail.id_nguoi_ban.ngay_cap_cmnd_cccd.strftime('%d/%m/%Y') if detail.id_nguoi_ban and detail.id_nguoi_ban.ngay_cap_cmnd_cccd else '',
                'dia_chi': detail.id_nguoi_ban.dia_chi if detail.id_nguoi_ban else '',
                'ngay_mua_hang': detail.ngay_mua_hang.strftime('%Y-%m-%d') if detail.ngay_mua_hang else date.today().strftime('%Y-%m-%d'),
                'ngay_mua_hang_display': detail.ngay_mua_hang.strftime('%d/%m/%Y') if detail.ngay_mua_hang else '',
                'ghi_chu': detail.ghi_chu or '',
                'ct_id': detail.id_ct_bang_ke_thu_mua_tu_dan 
                
            }
            formatted_details.append(formatted_detail)
        
        # Calculate totals
        tong_so_luong = sum(detail['so_luong'] for detail in formatted_details)
        tong_thanh_tien = sum(detail['thanh_tien'] for detail in formatted_details)
        
        # Get rollback info to display minimum production quantity
        rollback_record = BangKeTruLuiNguyenLieu.objects.filter(
            id_lenh_san_xuat=record.id_lenh_san_xuat,
            id_san_pham=record.id_san_pham
        ).first()
        so_luong_san_pham_xuat = rollback_record.so_luong_san_pham_xuat if rollback_record else 0

        so_luong_san_xuat_toi_thieu = int(so_luong_san_pham_xuat / ty_le_thu_hoi) if ty_le_thu_hoi > 0 else 0
        
        formatted_record = {
            'id_bang_ke_thu_mua_tu_dan': record.id_bang_ke_thu_mua_tu_dan,
            'ma_lenh_sx': record.id_lenh_san_xuat.id_lenh_san_xuat,
            'ma_don_hang': record.id_lenh_san_xuat.id_don_hang,
            'ten_nguyen_lieu': ten_nguyen_lieu,
            'ngay_lap_giay_to': record.ngay_lap_giay_to if record.ngay_lap_giay_to else '',
            'so_luong_san_pham_xuat': so_luong_san_pham_xuat,
            'ty_le_thu_hoi': f"{int(ty_le_thu_hoi * 100)}%" if ty_le_thu_hoi else "0%",
            'so_luong_san_xuat_toi_thieu': so_luong_san_xuat_toi_thieu,
            'chi_tiet': formatted_details,
            'tong_so_luong': tong_so_luong,
            'tong_thanh_tien': tong_thanh_tien,
            'ngay_from': ngay_from,
            'ngay_to': ngay_to,
            'don_vi_tinh': don_vi_tinh,
        }
        
        # Get nguoi list for dropdown
        nguoi_list = Nguoi.objects.all().values(
            'id', 'ten', 'so_cmnd_cccd', 'ngay_cap_cmnd_cccd', 'dia_chi', 'vai_tro'
        )
        
        context = {
            'record': formatted_record,
            'nguoi_list': list(nguoi_list)
        }
        
    except BangKeThuMuaTuDan.DoesNotExist:
        context = {'record': None, 'nguoi_list': []}
    
    return render(request, 'purchase_detail.html', context)


def handle_purchase_detail_update(request, record):
    """Handle purchase detail update"""
    try:
        # Update main record date range if provided
        ngay_tu_raw = request.POST.get('ngay_from', '')
        ngay_den_raw = request.POST.get('ngay_to', '')
        
        if ngay_tu_raw and ngay_den_raw:
            ngay_tu_fmt = datetime.strptime(ngay_tu_raw, "%Y-%m-%d").strftime("%d/%m/%Y")
            ngay_den_fmt = datetime.strptime(ngay_den_raw, "%Y-%m-%d").strftime("%d/%m/%Y")
            ngay_lap_giay_to = f"{ngay_tu_fmt} - {ngay_den_fmt}"
            record.ngay_lap_giay_to = ngay_lap_giay_to
            record.save()
        
        # Get material name from VatTu
        ten_nguyen_lieu, *_ = get_vat_tu_info_general(record.id_san_pham)
        
        # Process detail updates and new records
        success = _process_purchase_detail_updates(request, record, ten_nguyen_lieu)
        
        if success:
            return JsonResponse({
                'status': 'success',
                'message': 'Cập nhật thành công'
            })
        else:
            return JsonResponse({
                'status': 'error',
                'message': 'Có lỗi xảy ra khi cập nhật'
            })
            
    except Exception as e:
        print(f"Error updating purchase detail: {e}")
        return JsonResponse({
            'status': 'error',
            'message': f'Lỗi khi cập nhật: {str(e)}'
        })


def _process_purchase_detail_updates(request, purchase_record, ten_nguyen_lieu):
    """Process purchase detail record updates and new additions"""
    try:
        # Parse POST data by row index
        form_data = defaultdict(dict)

        for key, value in request.POST.items():
            if '_' not in key:
                continue
            parts = key.rsplit('_', 1)
            if len(parts) != 2 or not parts[1].isdigit():
                continue

            field, idx = parts
            form_data[idx][field] = value.strip()

        # Get existing detail records from DB
        existing_details = {
            str(detail.id_ct_bang_ke_thu_mua_tu_dan): detail
            for detail in CtBangKeThuMuaTuDan.objects.filter(id_bang_ke_thu_mua_tu_dan=purchase_record)
        }

        processed_ids = set()

        # Process each row in form data
        for idx, fields in form_data.items():
            ct_id = fields.get("ct_id", "").strip()
            id_nguoi_ban = fields.get("ten_nguoi_ban", "").strip()
            so_luong = _safe_float_conversion(fields.get("so_luong", "0"))
            don_gia = _safe_float_conversion(fields.get("don_gia", "0"))
            ngay_mua_hang_str = fields.get("ngay_mua_hang", "")
            ghi_chu = fields.get("ghi_chu", "").strip()

            if not id_nguoi_ban:
                continue  # Người bán là bắt buộc

            so_luong = max(so_luong, 0)
            don_gia = max(don_gia, 0)

            try:
                nguoi_ban = Nguoi.objects.get(id=int(id_nguoi_ban))
            except (Nguoi.DoesNotExist, ValueError):
                continue

            ngay_mua_hang = _parse_date(ngay_mua_hang_str) if ngay_mua_hang_str else date.today()

            if not ct_id:
                existing_similar = CtBangKeThuMuaTuDan.objects.filter(
                    id_bang_ke_thu_mua_tu_dan=purchase_record,
                    id_nguoi_ban=nguoi_ban,
                    ngay_mua_hang=ngay_mua_hang
                ).filter(
                    don_gia=don_gia,
                    so_luong=so_luong
                ).exclude(
                    id_ct_bang_ke_thu_mua_tu_dan__in=processed_ids
                ).exists()
                
                if existing_similar:
                    print(f"Duplicate row skipped: {nguoi_ban} - {ngay_mua_hang} - {so_luong} - {don_gia}")
                    continue

            if ct_id and ct_id in existing_details:
                detail_record = existing_details[ct_id]
            else:
                detail_record = CtBangKeThuMuaTuDan(
                    id_bang_ke_thu_mua_tu_dan=purchase_record
                )

            detail_record.ten_nguyen_lieu = ten_nguyen_lieu
            detail_record.don_gia = don_gia
            detail_record.so_luong = so_luong
            detail_record.id_nguoi_ban = nguoi_ban
            detail_record.ngay_mua_hang = ngay_mua_hang
            detail_record.ghi_chu = ghi_chu
            detail_record.save()

            if ct_id:
                processed_ids.add(ct_id)

        # Delete detail rows that were removed from the form
        for record_id, detail_record in existing_details.items():
            if record_id not in processed_ids:
                detail_record.delete()

        return True

    except Exception as e:
        print(f"Error in _process_purchase_detail_updates: {e}")
        return False

@csrf_exempt
@require_http_methods(["DELETE"])
def purchase_detail_row_delete(request, pk, row_id):
    """Delete a specific detail row"""
    try:
        # Get the main record first to verify access
        main_record = BangKeThuMuaTuDan.objects.get(pk=pk)
        
        # Get and delete the specific detail record
        detail_record = CtBangKeThuMuaTuDan.objects.get(
            id_ct_bang_ke_thu_mua_tu_dan=row_id,
            id_bang_ke_thu_mua_tu_dan=main_record
        )
        
        detail_record.delete()
        
        return JsonResponse({
            'status': 'success',
            'message': 'Đã xóa dòng thành công'
        })
        
    except (BangKeThuMuaTuDan.DoesNotExist, CtBangKeThuMuaTuDan.DoesNotExist):
        return JsonResponse({
            'status': 'error',
            'message': 'Không tìm thấy dữ liệu cần xóa'
        })
    except Exception as e:
        print(f"Error deleting detail row {row_id}: {e}")
        return JsonResponse({
            'status': 'error',
            'message': f'Lỗi khi xóa: {str(e)}'
        })

@csrf_exempt
@require_http_methods(["DELETE"])
def purchase_delete(request, pk):
    try:
        record = BangKeThuMuaTuDan.objects.get(pk=pk)
        
        # Lấy tên nguyên liệu từ bảng VatTu
        try:
            vat_tu = VatTu.objects.get(id_san_pham=record.id_san_pham)
            ten_nguyen_lieu = vat_tu.ten_khac or vat_tu.ten_sp_chinh
        except VatTu.DoesNotExist:
            ten_nguyen_lieu = f"Sản phẩm {record.id_san_pham}"
        
        # Lưu thông tin để log
        ma_lenh_sx = record.id_lenh_san_xuat.id_lenh_san_xuat
        
        # Xóa các bản ghi chi tiết trước (CASCADE delete)
        detail_count = CtBangKeThuMuaTuDan.objects.filter(
            id_bang_ke_thu_mua_tu_dan=record
        ).count()
        
        # Xóa bản ghi chính (sẽ tự động xóa các chi tiết nếu có CASCADE)
        record.delete()

        return JsonResponse({
            'status': 'success',
            'message': f'Đã xóa bảng kê "{ten_nguyen_lieu}" thành công'
        })
    except BangKeThuMuaTuDan.DoesNotExist:
        return JsonResponse({
            'status': 'error',
            'message': 'Không tìm thấy bảng kê cần xóa'
        })
    except Exception as e:
        print(f"Error deleting purchase record {pk}: {e}")
        return JsonResponse({
            'status': 'error',
            'message': f'Lỗi khi xóa: {str(e)}'
        })


def purchase_create(request):
    """Create new purchase from people records"""
    don_hang_list = LenhSanXuat.objects.values_list('id_don_hang', flat=True).distinct()
    
    if request.method == 'GET':
        return render(request, 'purchase_create.html', {'don_hang_list': don_hang_list})
    
    # POST request processing
    ma_lenh_sx = request.POST.get('ma_lenh_sx')
    id_san_pham = request.POST.get('ten_nguyen_lieu')  # Đây là id_san_pham được chọn
    
    if not ma_lenh_sx or not id_san_pham:
        return render(request, 'purchase_create.html', {
            'error': 'Vui lòng chọn lệnh sản xuất và nguyên liệu',
            'don_hang_list': don_hang_list
        })
    
    # Get production order
    try:
        lenh_sx = LenhSanXuat.objects.get(id_lenh_san_xuat=ma_lenh_sx)
    except LenhSanXuat.DoesNotExist:
        return render(request, 'purchase_create.html', {
            'error': f'Không tìm thấy lệnh sản xuất với mã: {ma_lenh_sx}',
            'don_hang_list': don_hang_list
        })
    
    # Check if purchase record already exists for this material and production order
    existing_record = BangKeThuMuaTuDan.objects.filter(
        id_lenh_san_xuat=lenh_sx,
        id_san_pham=id_san_pham
    ).first()
    
    if existing_record:
        return render(request, 'purchase_create.html', {
            'error': 'Đã tồn tại bảng kê thu mua cho nguyên liệu này',
            'don_hang_list': don_hang_list
        })
    ngay_tu_raw = request.POST.get('ngay_from', '')
    ngay_den_raw = request.POST.get('ngay_to', '')

    ngay_tu_fmt = datetime.strptime(ngay_tu_raw, "%Y-%m-%d").strftime("%d/%m/%Y") if ngay_tu_raw else ''
    ngay_den_fmt = datetime.strptime(ngay_den_raw, "%Y-%m-%d").strftime("%d/%m/%Y") if ngay_den_raw else ''

    ngay_lap_giay_to = f"{ngay_tu_fmt} - {ngay_den_fmt}" if ngay_tu_fmt and ngay_den_fmt else ''

    # Create main purchase record
    purchase_record = BangKeThuMuaTuDan(
        id_lenh_san_xuat=lenh_sx,
        id_san_pham=id_san_pham,
        ngay_lap_giay_to=ngay_lap_giay_to
    )
    purchase_record.save()
    
    # Process purchase details
    success = _process_purchase_details(request, purchase_record)
    
    if success:
        return redirect('purchase')
    else:
        # Delete the main record if detail processing failed
        purchase_record.delete()
        return render(request, 'purchase_create.html', {
            'error': 'Không thể tạo bảng kê thu mua. Vui lòng kiểm tra lại dữ liệu.',
            'don_hang_list': don_hang_list
        })


def _process_purchase_details(request, purchase_record):
    """Process purchase detail records"""
    # Get form data arrays
    ho_ten_list = request.POST.getlist('ho_ten[]')
    ngay_mua_hang_list = request.POST.getlist('ngay_mua_hang[]')
    so_luong_list = request.POST.getlist('so_luong[]')
    don_gia_list = request.POST.getlist('don_gia[]')
    ghi_chu_list = request.POST.getlist('ghi_chu[]')
    
    if not ho_ten_list:
        return False
    
    # Get material name from VatTu
    ten_nguyen_lieu, *_ = get_vat_tu_info_general(purchase_record.id_san_pham)

    # Process each detail record
    detail_created = False
    for i in range(len(ho_ten_list)):
        try:
            # Parse seller ID
            id_nguoi_ban = None
            if i < len(ho_ten_list) and ho_ten_list[i]:
                try:
                    id_nguoi_ban = Nguoi.objects.get(id=int(ho_ten_list[i]))
                except (Nguoi.DoesNotExist, ValueError):
                    continue  # Skip invalid seller
            
            # Parse other fields
            so_luong = _safe_float_conversion(so_luong_list[i] if i < len(so_luong_list) else None)
            don_gia = _safe_float_conversion(don_gia_list[i] if i < len(don_gia_list) else None)
            
            # Parse purchase date
            ngay_mua_hang = None
            if i < len(ngay_mua_hang_list) and ngay_mua_hang_list[i]:
                ngay_mua_hang = _parse_date(ngay_mua_hang_list[i])
            else:
                ngay_mua_hang = date.today()
            
            ghi_chu = ghi_chu_list[i] if i < len(ghi_chu_list) else ''
            
            # Validate required fields
            if not id_nguoi_ban or so_luong <= 0 or don_gia <= 0:
                continue
            
            # Create detail record
            detail_record = CtBangKeThuMuaTuDan(
                id_bang_ke_thu_mua_tu_dan=purchase_record,
                ten_nguyen_lieu=ten_nguyen_lieu,
                don_gia=don_gia,
                so_luong=so_luong,
                id_nguoi_ban=id_nguoi_ban,
                ngay_mua_hang=ngay_mua_hang,
                ghi_chu=ghi_chu
            )
            detail_record.save()
            detail_created = True
                
        except Exception as e:
            print(f"Error processing purchase detail {i}: {e}")
            continue
    
    return detail_created


@require_GET
def get_lenh_sx_by_don_hang(request, ma_don_hang):
    """Trả danh sách lệnh sản xuất theo mã đơn hàng"""
    try:
        lenh_sx_list = LenhSanXuat.objects.filter(
            id_don_hang=ma_don_hang
        ).values_list('id_lenh_san_xuat', flat=True)
        
        return JsonResponse({
            'status': 'success',
            'data': list(lenh_sx_list)
        })
    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': f'Lỗi: {str(e)}'
        })


@require_GET  
def get_purchase_materials_by_lenh_sx_api(request, ma_lenh_sx):
    """API endpoint để lấy materials theo mã lệnh sản xuất cho purchase create"""
    try:
        lenh_sx = LenhSanXuat.objects.get(id_lenh_san_xuat=ma_lenh_sx)
        
        # Get all materials from rollback records for this production order (chỉ lấy NVL - KHÔ)
        rollback_materials = BangKeTruLuiNguyenLieu.objects.filter(
            id_lenh_san_xuat=lenh_sx,
            id_san_pham__in=VatTu.objects.filter(nhom_vthh='NVL - KHÔ').values_list('id_san_pham', flat=True)
        )
        
        # Get materials that already have purchase records
        existing_purchase_materials = BangKeThuMuaTuDan.objects.filter(
            id_lenh_san_xuat=lenh_sx
        ).values_list('id_san_pham', flat=True)
        
        # Filter out materials that already have purchase records (chỉ hiện thị chưa tạo)
        available_materials = rollback_materials.exclude(
            id_san_pham__in=existing_purchase_materials
        )
        
        materials_data = []
        for material in available_materials:
            # Get VatTu info
            try:
                vat_tu = VatTu.objects.get(id_san_pham=material.id_san_pham, nhom_vthh='NVL - KHÔ')
                ten_nguyen_lieu = vat_tu.ten_khac or vat_tu.ten_sp_chinh
                ty_le_thu_hoi = vat_tu.ty_le_thu_hoi or 0
            except VatTu.DoesNotExist:
                ten_nguyen_lieu = material.ten_nguyen_lieu or f"Sản phẩm {material.id_san_pham}"
                ty_le_thu_hoi = 0
            
            # Get production order output quantity từ bảng kê trừ lùi
            so_luong_san_pham_xuat = material.so_luong_san_pham_xuat or 0
            
            # Calculate minimum production quantity = số lượng sản phẩm xuất / tỷ lệ thu hồi
            so_luong_san_xuat_toi_thieu = 0
            if ty_le_thu_hoi > 0 and so_luong_san_pham_xuat:
                so_luong_san_xuat_toi_thieu = so_luong_san_pham_xuat / ty_le_thu_hoi
            
            material_data = {
                'id_san_pham': material.id_san_pham,
                'ten_nguyen_lieu': ten_nguyen_lieu,
                'so_luong_san_pham_xuat': so_luong_san_pham_xuat,
                'ty_le_thu_hoi': ty_le_thu_hoi,
                'ty_le_thu_hoi_percent': f"{int(ty_le_thu_hoi * 100)}%" if ty_le_thu_hoi else "0%",
                'so_luong_san_xuat_toi_thieu': round(so_luong_san_xuat_toi_thieu, 2) if so_luong_san_xuat_toi_thieu else 0
            }
            materials_data.append(material_data)
        
        return JsonResponse({
            'status': 'success',
            'materials': materials_data
        })
        
    except LenhSanXuat.DoesNotExist:
        return JsonResponse({
            'status': 'error',
            'message': f'Không tìm thấy lệnh sản xuất với mã: {ma_lenh_sx}'
        })
    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': f'Lỗi: {str(e)}'
        })


@require_GET
def get_nguoi_list(request):
    """API endpoint to get list of people for dropdown"""
    try:
        nguoi_list = Nguoi.objects.all().values(
            'id', 'ten', 'so_cmnd_cccd', 'ngay_cap_cmnd_cccd', 'dia_chi', 'vai_tro'
        )
        
        return JsonResponse({
            'status': 'success',
            'nguoi_list': list(nguoi_list)
        })
        
    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': f'Lỗi khi lấy danh sách người: {str(e)}'
        })


@require_GET
def get_nguoi_detail(request):
    """API endpoint to get person details by ID"""
    nguoi_id = request.GET.get('id')
    
    if not nguoi_id:
        return JsonResponse({
            'status': 'error',
            'message': 'Vui lòng cung cấp ID người'
        })
    
    try:
        nguoi = Nguoi.objects.get(id=nguoi_id)
        
        return JsonResponse({
            'status': 'success',
            'nguoi': {
                'id': nguoi.id,
                'ten': nguoi.ten,
                'so_cmnd_cccd': nguoi.so_cmnd_cccd or '',
                'ngay_cap_cmnd_cccd': nguoi.ngay_cap_cmnd_cccd.strftime('%d/%m/%Y') if nguoi.ngay_cap_cmnd_cccd else '',
                'dia_chi': nguoi.dia_chi or '',
                'vai_tro': nguoi.vai_tro or ''
            }
        })
        
    except Nguoi.DoesNotExist:
        return JsonResponse({
            'status': 'error',
            'message': f'Không tìm thấy người với ID: {nguoi_id}'
        })

@csrf_exempt
@require_POST
def add_nguoi(request):
    """API endpoint to add new person"""
    try:
        data = json.loads(request.body)
        
        ten = data.get('ten', '').strip()
        so_cmnd_cccd = data.get('so_cmnd_cccd', '').strip()
        ngay_cap_cmnd_cccd_str = data.get('ngay_cap_cmnd_cccd', '').strip()
        ngay_cap_cmnd_cccd = _parse_date(ngay_cap_cmnd_cccd_str) if ngay_cap_cmnd_cccd_str else None
        dia_chi = data.get('dia_chi', '').strip()
        vai_tro = data.get('vai_tro', 'Người bán').strip()
        
        if not ten:
            return JsonResponse({
                'status': 'error',
                'message': 'Vui lòng nhập tên người'
            })
        
        # Create new person
        nguoi = Nguoi(
            ten=ten,
            so_cmnd_cccd=so_cmnd_cccd if so_cmnd_cccd else None,
            ngay_cap_cmnd_cccd=ngay_cap_cmnd_cccd,
            dia_chi=dia_chi if dia_chi else None,
            vai_tro=vai_tro
        )
        nguoi.save()
        
        return JsonResponse({
            'status': 'success',
            'message': 'Thêm người thành công',
            'nguoi': {
                'id': nguoi.id,
                'ten': nguoi.ten,
                'so_cmnd_cccd': nguoi.so_cmnd_cccd or '',
                'ngay_cap_cmnd_cccd': nguoi.ngay_cap_cmnd_cccd.strftime('%d/%m/%Y') if nguoi.ngay_cap_cmnd_cccd else '',
                'dia_chi': nguoi.dia_chi or '',
                'vai_tro': nguoi.vai_tro or ''
            }
        })
        
    except json.JSONDecodeError:
        return JsonResponse({
            'status': 'error',
            'message': 'Dữ liệu JSON không hợp lệ'
        })
    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': f'Lỗi khi thêm người: {str(e)}'
        })

def _get_production_order_output_quantity(material):
    """Get production order output quantity from rollback record"""
    return material.so_luong_san_pham_xuat or 0

@require_GET
def purchase_export_pdf(request, pk):
    """Render purchase record as HTML for PDF printing (AJAX)"""
    try:
        record = BangKeThuMuaTuDan.objects.select_related('id_lenh_san_xuat').get(pk=pk)
    except BangKeThuMuaTuDan.DoesNotExist:
        return HttpResponse("Không tìm thấy bảng kê thu mua", status=404)

    # Get detail records
    detail_records = CtBangKeThuMuaTuDan.objects.filter(
        id_bang_ke_thu_mua_tu_dan=record
    ).select_related('id_nguoi_ban').order_by('ngay_mua_hang', 'id_ct_bang_ke_thu_mua_tu_dan')
    
    # Get VatTu info
    try:
        vat_tu = VatTu.objects.get(id_san_pham=record.id_san_pham)
        ten_nguyen_lieu = vat_tu.ten_khac or vat_tu.ten_sp_chinh
        don_vi_tinh = vat_tu.don_vi_tinh or 'Kg'
    except VatTu.DoesNotExist:
        ten_nguyen_lieu = f"Sản phẩm {record.id_san_pham}"
        don_vi_tinh = 'Kg'

    # Format detail records for display
    formatted_details = []
    for detail in detail_records:
        formatted_detail = {
            'ngay_mua_hang': detail.ngay_mua_hang,
            'ngay_mua_hang_display': detail.ngay_mua_hang.strftime('%d/%m/%Y') if detail.ngay_mua_hang else '',
            'ten_nguoi_ban': detail.id_nguoi_ban.ten if detail.id_nguoi_ban else '',
            'so_cmnd_cccd': detail.id_nguoi_ban.so_cmnd_cccd if detail.id_nguoi_ban else '',
            'dia_chi': detail.id_nguoi_ban.dia_chi if detail.id_nguoi_ban else '',
            'ten_hang': ten_nguyen_lieu,
            'don_vi_tinh': don_vi_tinh,
            'so_luong': detail.so_luong or 0,
            'don_gia': detail.don_gia or 0,
            'thanh_tien': (detail.don_gia or 0) * (detail.so_luong or 0),
            'ghi_chu': detail.ghi_chu or ''
        }
        formatted_details.append(formatted_detail)

    # Calculate totals
    tong_so_luong = sum(detail['so_luong'] for detail in formatted_details)
    tong_thanh_tien = sum(detail['thanh_tien'] for detail in formatted_details)
    
    # Convert total amount to Vietnamese words
    tong_thanh_tien_chu = convert_number_to_vietnamese_words(int(tong_thanh_tien))
    
    # Parse date range
    ngay_from = ''
    ngay_to = ''
    if record.ngay_lap_giay_to and ' - ' in record.ngay_lap_giay_to:
        parts = record.ngay_lap_giay_to.split(' - ')
        ngay_from = parts[0] if len(parts) > 0 else ''
        ngay_to = parts[1] if len(parts) > 1 else ''

    formatted_record = {
        'ma_lenh_sx': record.id_lenh_san_xuat.id_lenh_san_xuat,
        'ma_don_hang': record.id_lenh_san_xuat.id_don_hang,
        'ten_nguyen_lieu': ten_nguyen_lieu,
        'ngay_from': ngay_from,
        'ngay_to': ngay_to,
        'chi_tiet': formatted_details,
        'tong_so_luong': tong_so_luong,
        'tong_thanh_tien': tong_thanh_tien,
        'tong_thanh_tien_chu': tong_thanh_tien_chu,
        'don_vi_tinh': don_vi_tinh,
        'ngay_in': datetime.now().date()
    }

    context = {
        'record': formatted_record
    }

    return render(request, 'purchase_detail_pdf.html', context)

@require_GET
def purchase_export_excel(request, pk):
    """Export purchase record to Excel file"""
    try:
        record = BangKeThuMuaTuDan.objects.select_related('id_lenh_san_xuat').get(pk=pk)
    except BangKeThuMuaTuDan.DoesNotExist:
        return HttpResponse("Không tìm thấy bảng kê thu mua", status=404)

    # Get detail records
    detail_records = CtBangKeThuMuaTuDan.objects.filter(
        id_bang_ke_thu_mua_tu_dan=record
    ).select_related('id_nguoi_ban').order_by('ngay_mua_hang', 'id_ct_bang_ke_thu_mua_tu_dan')
    
    # Get VatTu info
    try:
        vat_tu = VatTu.objects.get(id_san_pham=record.id_san_pham)
        ten_nguyen_lieu = vat_tu.ten_khac or vat_tu.ten_sp_chinh
        don_vi_tinh = vat_tu.don_vi_tinh or 'Kg'
    except VatTu.DoesNotExist:
        ten_nguyen_lieu = f"Sản phẩm {record.id_san_pham}"
        don_vi_tinh = 'Kg'
    
    output = BytesIO()
    workbook = xlsxwriter.Workbook(output, {'in_memory': True})
    worksheet = workbook.add_worksheet("Bảng kê thu mua")

    # Define formats
    header_format = workbook.add_format({
        'font_name': 'Times New Roman',
        'font_size': 12,
        'bold': True,
        'align': 'left',
        'valign': 'vcenter'
    })
    
    title_format = workbook.add_format({
        'font_name': 'Times New Roman',
        'font_size': 14,
        'bold': True,
        'align': 'center',
        'valign': 'vcenter'
    })
    
    subtitle_format = workbook.add_format({
        'font_name': 'Times New Roman',
        'font_size': 14,
        'bold': True,
        'align': 'center',
        'valign': 'vcenter'
    })
    
    table_header_format = workbook.add_format({
        'font_name': 'Times New Roman',
        'font_size': 12,
        'bold': True,
        'align': 'center',
        'valign': 'vcenter',
        'border': 1,
        'bg_color': '#F2F2F2'
    })
    
    cell_format = workbook.add_format({
        'font_name': 'Times New Roman',
        'font_size': 12,
        'align': 'center',
        'valign': 'vcenter',
        'border': 1
    })
    
    cell_left_format = workbook.add_format({
        'font_name': 'Times New Roman',
        'font_size': 12,
        'align': 'left',
        'valign': 'vcenter',
        'border': 1
    })
    
    cell_right_format = workbook.add_format({
        'font_name': 'Times New Roman',
        'font_size': 12,
        'align': 'right',
        'valign': 'vcenter',
        'border': 1,
        'num_format': '#,##0'
    })
    
    total_format = workbook.add_format({
        'font_name': 'Times New Roman',
        'font_size': 12,
        'bold': True,
        'align': 'right',
        'valign': 'vcenter',
        'border': 1,
        'num_format': '#,##0'
    })

    # Set column widths
    worksheet.set_column('A:A', 4)   # STT
    worksheet.set_column('B:B', 10)  # Ngày
    worksheet.set_column('C:C', 18)  # Họ tên
    worksheet.set_column('D:D', 12)  # CMND
    worksheet.set_column('E:E', 25)  # Địa chỉ
    worksheet.set_column('F:F', 15)  # Tên hàng
    worksheet.set_column('G:G', 8)   # Đơn vị
    worksheet.set_column('H:H', 10)  # Số lượng
    worksheet.set_column('I:I', 12)  # Đơn giá
    worksheet.set_column('J:J', 15)  # Tổng tiền
    worksheet.set_column('K:K', 12)  # Ghi chú

    # Company header (left side)
    worksheet.write('A1', 'CÔNG TY CP TÂN PHONG', header_format)
    worksheet.write('A2', 'ĐC: TT Hùng Sơn - L.Thao - Phú Thọ', header_format)
    worksheet.write('A3', 'ĐT: 0210.22152277/38629388', header_format)

    # Title and material name (right side)
    worksheet.merge_range('E1:I2', f'BẢNG TỔNG HỢP THU MUA {ten_nguyen_lieu.upper()}', title_format)
    
    # Date range
    date_range = ''
    if record.ngay_lap_giay_to and ' - ' in record.ngay_lap_giay_to:
        date_range = f"Ngày {record.ngay_lap_giay_to}"
    
    worksheet.merge_range('E3:I3', date_range, subtitle_format)

    # Table headers
    row = 5
    worksheet.merge_range(row, 2, row, 4, 'Người bán', table_header_format)  # Gộp cột C, D, E
    worksheet.write(row, 0, 'STT', table_header_format)
    worksheet.write(row, 1, 'Ngày,tháng', table_header_format)
    worksheet.write(row, 5, 'Tên hàng', table_header_format)
    worksheet.write(row, 6, 'Đơn vị tính', table_header_format)
    worksheet.write(row, 7, 'Số lượng', table_header_format)
    worksheet.write(row, 8, 'Đơn giá', table_header_format)
    worksheet.write(row, 9, 'Tổng tiền', table_header_format)
    worksheet.write(row, 10, 'Ghi chú', table_header_format)
    row += 1
    worksheet.write(row, 0, '', table_header_format)  # STT trống
    worksheet.write(row, 1, '', table_header_format)  # Ngày trống
    worksheet.write(row, 2, 'Họ và tên', table_header_format)
    worksheet.write(row, 3, 'Số CMND/CCCD', table_header_format)
    worksheet.write(row, 4, 'Địa chỉ', table_header_format)
    worksheet.write(row, 5, '', table_header_format)  # Tên hàng trống
    worksheet.write(row, 6, '', table_header_format)  # Đơn vị trống
    worksheet.write(row, 7, '', table_header_format)  # Số lượng trống
    worksheet.write(row, 8, '', table_header_format)  # Đơn giá trống
    worksheet.write(row, 9, '', table_header_format)  # Tổng tiền trống
    worksheet.write(row, 10, '', table_header_format)  # Ghi chú trống
    # Data rows
    row += 1
    start_data_row = row
    for idx, detail in enumerate(detail_records, 1):
        worksheet.write(row, 0, idx, cell_format)
        worksheet.write(row, 1, detail.ngay_mua_hang.strftime('%d/%m/%Y') if detail.ngay_mua_hang else '', cell_format)
        worksheet.write(row, 2, detail.id_nguoi_ban.ten if detail.id_nguoi_ban else '', cell_left_format)
        worksheet.write(row, 3, detail.id_nguoi_ban.so_cmnd_cccd if detail.id_nguoi_ban else '', cell_format)
        worksheet.write(row, 4, detail.id_nguoi_ban.dia_chi if detail.id_nguoi_ban else '', cell_left_format)
        worksheet.write(row, 5, ten_nguyen_lieu, cell_left_format)
        worksheet.write(row, 6, don_vi_tinh, cell_format)
        worksheet.write(row, 7, detail.so_luong or 0, cell_right_format)
        worksheet.write(row, 8, detail.don_gia or 0, cell_right_format)
        worksheet.write(row, 9, (detail.don_gia or 0) * (detail.so_luong or 0), cell_right_format)
        worksheet.write(row, 10, detail.ghi_chu or '', cell_left_format)
        row += 1

    # Totals row TRONG bảng (không phải ngoài bảng)
    tong_so_luong = sum((detail.so_luong or 0) for detail in detail_records)
    tong_thanh_tien = sum((detail.don_gia or 0) * (detail.so_luong or 0) for detail in detail_records)

    # Dòng tổng trong bảng (chỉ hiển thị số liệu ở cột số lượng và thành tiền)
    worksheet.write(row, 0, '', cell_format)  # STT trống
    worksheet.write(row, 1, '', cell_format)  # Ngày trống  
    worksheet.write(row, 2, '', cell_format)  # Họ tên trống
    worksheet.write(row, 3, '', cell_format)  # CMND trống
    worksheet.write(row, 4, '', cell_format)  # Địa chỉ trống
    worksheet.write(row, 5, '', cell_format)  # Tên hàng trống
    worksheet.write(row, 6, '', cell_format)  # Đơn vị trống
    worksheet.write(row, 7, tong_so_luong, total_format)  # Tổng số lượng
    worksheet.write(row, 8, '', cell_format)  # Đơn giá trống
    worksheet.write(row, 9, tong_thanh_tien, total_format)  # Tổng tiền
    worksheet.write(row, 10, '', cell_format)  # Ghi chú trống

    # Tổng cộng (ngoài bảng)
    row += 1
    worksheet.write(row, 2, f'Tổng cộng: {tong_so_luong:,.0f}Kg = {tong_thanh_tien:,.0f} Vnđ', 
                workbook.add_format({'font_name': 'Times New Roman', 'font_size': 12, 'bold': True}))

    # Bằng chữ (ngoài bảng, thẳng cột với "Họ và tên")
    row += 1
    tong_thanh_tien_chu = convert_number_to_vietnamese_words(int(tong_thanh_tien))
    worksheet.write(row, 2, f'Bằng chữ: {tong_thanh_tien_chu} đồng chẵn.', 
                workbook.add_format({'font_name': 'Times New Roman', 'font_size': 12, 'italic': True}))

    # Signatures (thẳng cột với "Họ và tên", "Địa chỉ", "Tổng tiền")
    row += 2
    signature_format = workbook.add_format({
        'font_name': 'Times New Roman',
        'font_size': 12,
        'bold': True,
        'align': 'center'
    })

    worksheet.write(row, 2, 'Người thu mua', signature_format)  # Cột C (Họ và tên)
    worksheet.write(row, 4, 'Kế toán', signature_format)        # Cột E (Địa chỉ) 
    worksheet.write(row, 8, 'CTHĐQT-Tổng giám đốc', signature_format)  # Cột J (Tổng tiền)

    # Signature names
    row += 5
    name_format = workbook.add_format({
        'font_name': 'Times New Roman',
        'font_size': 12,
        'bold': True,
        'align': 'center'
    })

    worksheet.write(row, 2, 'Đinh Thị Thanh', name_format)
    worksheet.write(row, 4, 'Phạm Thị Thu Hương', name_format)
    worksheet.write(row, 8, 'Đinh Thị Thức', name_format)  # Thêm tên Tổng giám đốc

    # Confirmation section
    row += 2
    worksheet.write(row + 1, 4, 'XÁC NHẬN CỦA ĐỊA PHƯƠNG NƠI THU MUA', 
                workbook.add_format({
                    'font_name': 'Times New Roman',
                    'font_size': 12,
                    'bold': True,
                    'align': 'center'
                }))
    workbook.close()
    output.seek(0)

    # Create HTTP response
    response = HttpResponse(
        output.read(),
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    
    filename = f"BangKeThuMua_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    
    return response

def convert_number_to_vietnamese_words(number):
    try:
        number = int(number)
    except Exception:
        return ""
    words = num2words(number, lang='vi')
    return words[0].upper() + words[1:] 


# ==================== WO MANAGEMENT ====================

def wo_ledger(request):
    """Main WO ledger view with filtering"""
    ma_don_hang = request.GET.get('ma_don_hang')
    ma_lenh_sx = request.GET.get('ma_lenh_sx')
    
    # Get all WO records with related data
    wo_records = BangKeWo.objects.select_related(
        'id_lenh_san_xuat',
        'id_nguoi'
    ).all()

    # Apply filters
    if ma_don_hang:
        wo_records = wo_records.filter(id_lenh_san_xuat__id_don_hang=ma_don_hang)
    if ma_lenh_sx:
        wo_records = wo_records.filter(id_lenh_san_xuat__id_lenh_san_xuat=ma_lenh_sx)

    formatted_records = []
    for wo_record in wo_records:
        # Get material info
        ten_nguyen_lieu, _, ma_hs, _ = get_vat_tu_info_general(wo_record.id_san_pham)

        # Get product output quantity from rollback table
        rollback_record = BangKeTruLuiNguyenLieu.objects.filter(
            id_lenh_san_xuat=wo_record.id_lenh_san_xuat,
            id_san_pham=wo_record.id_san_pham
        ).first()
        so_luong_san_pham_xuat = rollback_record.so_luong_san_pham_xuat if rollback_record else 0

        formatted_record = {
            'id': wo_record.id,
            'ma_lenh_sx': wo_record.id_lenh_san_xuat.id_lenh_san_xuat,
            'ten_nguyen_lieu': ten_nguyen_lieu,
            'ma_hs': ma_hs,
            'so_luong_san_pham_xuat': so_luong_san_pham_xuat,
            'tri_gia_fob': f"{wo_record.tri_gia_fob:,.0f}" if wo_record.tri_gia_fob else "0",
        }
        formatted_records.append(formatted_record)

    # Get distinct order list for filter
    don_hang_list = LenhSanXuat.objects.values_list('id_don_hang', flat=True).distinct()

    context = {
        'records': formatted_records,
        'don_hang_list': don_hang_list
    }

    return render(request, 'wo_management.html', context)


def wo_ledger_detail(request, pk):
    """WO ledger detail view with edit functionality - SỬA CHỮA"""
    try:
        wo_record = BangKeWo.objects.select_related('id_lenh_san_xuat', 'id_nguoi').get(pk=pk)
        
        if request.method == 'POST':
            # Handle update request
            return handle_wo_detail_update(request, wo_record)
        
        # GET request - display detail
        lenh_san_xuat = wo_record.id_lenh_san_xuat

        # Get material info từ VatTu (tên nguyên liệu gốc - readonly)
        ten_nguyen_lieu_goc, don_vi_tinh_goc, ma_hs, _ = get_vat_tu_info_general(wo_record.id_san_pham)

        # Get rollback info
        rollback_record = BangKeTruLuiNguyenLieu.objects.filter(
            id_lenh_san_xuat=lenh_san_xuat,
            id_san_pham=wo_record.id_san_pham
        ).first()
        so_luong_san_pham_xuat = rollback_record.so_luong_san_pham_xuat if rollback_record else 0

        # Get purchase record info (for readonly fields)
        try:
            purchase_record = BangKeThuMuaTuDan.objects.get(
                id_bang_ke_thu_mua_tu_dan=wo_record.id_bang_ke_thu_mua_tu_dan
            )
            try:
                vat_tu_purchase = VatTu.objects.get(id_san_pham=purchase_record.id_san_pham)
                ten_hang_hoa_goc = vat_tu_purchase.ten_khac or ''
            except:
                ten_hang_hoa_goc = f"Sản phẩm {purchase_record.id_san_pham}"
        except:
            purchase_record = None
            ten_hang_hoa_goc = ten_nguyen_lieu_goc

        # Format record data - LOGIC GIỐNG PHẦN TẠO MỚI
        formatted_record = {
            'id': wo_record.id,
            'id_bang_ke_thu_mua_tu_dan': wo_record.id_bang_ke_thu_mua_tu_dan,
            'ma_lenh_sx': lenh_san_xuat.id_lenh_san_xuat,
            'ma_don_hang': lenh_san_xuat.id_don_hang,
            
            # Readonly fields từ VatTu
            'ten_nguyen_lieu_goc': ten_nguyen_lieu_goc,  # Từ VatTu (readonly)
            'ma_hs': ma_hs,  # Từ VatTu (readonly)
            'so_luong_san_pham_xuat': so_luong_san_pham_xuat,  # Từ bảng trừ lùi (readonly)
            'ten_hang_hoa_goc': ten_hang_hoa_goc,  # Tên hàng hóa từ bảng kê thu mua (readonly)
            
            # Editable fields - WO specific
            'ten_nguyen_lieu': wo_record.ten_hang_hoa or ten_nguyen_lieu_goc,  # Tên nguyên liệu WO (editable)
            'don_vi_tinh': don_vi_tinh_goc,  # Đơn vị tính WO (editable)
            'tri_gia_fob': wo_record.tri_gia_fob or 0,
            'to_khai_hai_quan': wo_record.to_khai_hai_quan or '',
            'dia_chi_thu_mua': wo_record.dia_chi_thu_mua or '',
            'noi_khai_thac': wo_record.noi_khai_thac or '',
            'so_luong_wo': wo_record.so_luong or 0,
            'ngay': wo_record.ngay.strftime('%Y-%m-%d') if wo_record.ngay else date.today().strftime('%Y-%m-%d'),
            'ngay_display': wo_record.ngay.strftime('%d/%m/%Y') if wo_record.ngay else '',
            'ghi_chu': wo_record.ghi_chu or '',
            'nguoi_phu_trach_id': wo_record.id_nguoi.id if wo_record.id_nguoi else None,
            'ten_nguoi': wo_record.id_nguoi.ten if wo_record.id_nguoi else '',
            'cccd_cmnd': wo_record.id_nguoi.so_cmnd_cccd if wo_record.id_nguoi else '',
        }
        
        # Get nguoi list for dropdown (Người mua)
        nguoi_list = Nguoi.objects.filter(vai_tro='Người mua').values(
            'id', 'ten', 'so_cmnd_cccd', 'ngay_cap_cmnd_cccd', 'dia_chi', 'vai_tro'
        )

        # Get purchase details for table display - LOGIC GIỐNG PHẦN TẠO MỚI
        purchase_details = []
        try:
            detail_records = CtBangKeThuMuaTuDan.objects.filter(
                id_bang_ke_thu_mua_tu_dan=wo_record.id_bang_ke_thu_mua_tu_dan
            ).select_related('id_nguoi_ban')
            
            for detail in detail_records:
                nguoi = detail.id_nguoi_ban
                
                # Format CCCD với ngày cấp
                cccd_info = ''
                if nguoi and nguoi.so_cmnd_cccd:
                    cccd_info = nguoi.so_cmnd_cccd
                    if nguoi.ngay_cap_cmnd_cccd:
                        cccd_info += f" cấp ngày {nguoi.ngay_cap_cmnd_cccd.strftime('%d/%m/%Y')}"
                
                purchase_details.append({
                    'ngay_mua_hang': detail.ngay_mua_hang.strftime('%d/%m/%Y') if detail.ngay_mua_hang else '',
                    'ten_nguoi_ban': nguoi.ten if nguoi else '',
                    'dia_chi': nguoi.dia_chi if nguoi else '',
                    'so_cmnd_cccd': cccd_info,
                    'so_luong': detail.so_luong or 0,
                    'don_gia': detail.don_gia or 0,  # Trả về số thay vì format
                    'thanh_tien': (detail.so_luong or 0) * (detail.don_gia or 0),
                    'ghi_chu': detail.ghi_chu or '',
                })
        except:
            pass
        
        context = {
            'record': formatted_record,
            'nguoi_list': list(nguoi_list),
            'purchase_details': purchase_details,
            'is_detail_view': True
        }
        
    except BangKeWo.DoesNotExist:
        context = {
            'record': None, 
            'wo_data': None,
            'nguoi_list': [],
            'purchase_details': [],
            'is_detail_view': True
        }
    
    return render(request, 'wo_detail.html', context)


def handle_wo_detail_update(request, wo_record):
    """Handle WO detail update - SỬA CHỮA MAPPING"""
    try:
        data = json.loads(request.body)
        
        # Update WO record fields
        wo_record.to_khai_hai_quan = data.get('to_khai_hai_quan', '')
        wo_record.dia_chi_thu_mua = data.get('dia_chi_thu_mua', '')
        wo_record.noi_khai_thac = data.get('noi_khai_thac', '')
        wo_record.so_luong = data.get('so_luong_wo', 0)
        wo_record.tri_gia_fob = data.get('tri_gia_fob', 0)
        wo_record.ten_hang_hoa = data.get('ten_nguyen_lieu', '')  # Map tên nguyên liệu WO
        wo_record.don_vi_tinh = data.get('don_vi_tinh', 'KGM')  # Thêm đơn vị tính WO
        wo_record.ghi_chu = data.get('ghi_chu', '')
        
        # Handle date
        ngay_str = data.get('ngay')
        if ngay_str:
            try:
                wo_record.ngay = datetime.strptime(ngay_str, '%Y-%m-%d').date()
            except:
                wo_record.ngay = date.today()
        else:
            wo_record.ngay = date.today()
        
        # Handle nguoi
        id_nguoi = data.get('id_nguoi')
        if id_nguoi:
            try:
                nguoi = Nguoi.objects.get(id=id_nguoi)
                wo_record.id_nguoi = nguoi
            except Nguoi.DoesNotExist:
                wo_record.id_nguoi = None
        else:
            wo_record.id_nguoi = None
            
        wo_record.save()

        return JsonResponse({
            'status': 'success',
            'message': 'Cập nhật bảng kê WO thành công',
            'data': {
                'id': wo_record.id,
                'tri_gia_fob': wo_record.tri_gia_fob or 0,
            }
        })

    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': f'Lỗi khi cập nhật: {str(e)}'
        }, status=400)

@require_http_methods(["DELETE"])
@csrf_exempt
def delete_wo_record(request, pk):
    """Delete WO record"""
    try:
        wo_record = get_object_or_404(BangKeWo, pk=pk)
        wo_record.delete()
        
        return JsonResponse({
            'status': 'success',
            'message': 'Xóa bảng kê WO thành công'
        })
    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': f'Lỗi khi xóa: {str(e)}'
        }, status=400)


def create_wo_record(request):
    """Create new WO record"""
    if request.method == 'GET':
        # Get distinct order list for initial dropdown
        don_hang_list = LenhSanXuat.objects.values_list('id_don_hang', flat=True).distinct().order_by('id_don_hang')
        
        # Get nguoi list with role filter
        nguoi_list = Nguoi.objects.filter(vai_tro='Người mua').values(
            'id', 'ten', 'so_cmnd_cccd', 'ngay_cap_cmnd_cccd', 'dia_chi', 'vai_tro'
        )
        
        context = {
            'don_hang_list': list(don_hang_list),
            'nguoi_list': list(nguoi_list)
        }
        
        return render(request, 'wo_create.html', context)
    
    elif request.method == 'POST':
        # Handle create new WO
        try:
            data = json.loads(request.body)
            id_bang_ke_thu_mua_tu_dan = data.get('id_bang_ke_thu_mua_tu_dan')
            
            if not id_bang_ke_thu_mua_tu_dan:
                return JsonResponse({
                    'status': 'error',
                    'message': 'Vui lòng chọn bảng kê thu mua từ dân'
                }, status=400)
            
            # Get purchase record
            purchase_record = get_object_or_404(BangKeThuMuaTuDan, pk=id_bang_ke_thu_mua_tu_dan)
            
            # Check if WO already exists
            if BangKeWo.objects.filter(id_bang_ke_thu_mua_tu_dan=id_bang_ke_thu_mua_tu_dan).exists():
                return JsonResponse({
                    'status': 'error',
                    'message': 'Bảng kê WO cho bản ghi này đã tồn tại'
                }, status=400)
            
            # Handle date
            ngay_str = data.get('ngay')
            ngay_obj = date.today()
            if ngay_str:
                try:
                    ngay_obj = datetime.strptime(ngay_str, '%Y-%m-%d').date()
                except:
                    pass
            
            # Handle nguoi
            nguoi_obj = None
            id_nguoi = data.get('id_nguoi')
            if id_nguoi:
                try:
                    nguoi_obj = Nguoi.objects.get(id=id_nguoi)
                except Nguoi.DoesNotExist:
                    pass
            
            # Create WO record
            wo_record = BangKeWo.objects.create(
                id_lenh_san_xuat=purchase_record.id_lenh_san_xuat,
                id_san_pham=purchase_record.id_san_pham,
                id_nguoi=nguoi_obj,
                id_bang_ke_thu_mua_tu_dan=id_bang_ke_thu_mua_tu_dan,
                to_khai_hai_quan=data.get('to_khai_hai_quan', ''),
                dia_chi_thu_mua=data.get('dia_chi_thu_mua', ''),
                noi_khai_thac=data.get('noi_khai_thac', ''),
                so_luong=data.get('so_luong_wo', 0),
                tri_gia_fob=data.get('tri_gia_fob', 0),
                ten_hang_hoa=data.get('ten_nguyen_lieu', ''),  # Tên hàng hóa từ bảng kê thu mua
                ngay=ngay_obj,
                #ghi_chu=data.get('ghi_chu', ''),
            )
            
            return JsonResponse({
                'status': 'success',
                'message': 'Tạo bảng kê WO thành công',
                'data': {
                    'id': wo_record.id,
                    'redirect_url': f'/wo/'
                }
            })
            
        except Exception as e:
            return JsonResponse({
                'status': 'error',
                'message': f'Lỗi khi tạo: {str(e)}'
            }, status=400)


# ==================== API ENDPOINTS ====================

def get_lenh_san_xuat_for_wo(request):
    """API endpoint to get production orders for WO filtering"""
    ma_don_hang = request.GET.get('ma_don_hang')
    
    if not ma_don_hang:
        return JsonResponse({'lenh_sx_list': []})
    
    try:
        # Get production orders that have WO records
        lenh_sx_with_wo = BangKeWo.objects.filter(
            id_lenh_san_xuat__id_don_hang=ma_don_hang
        ).values_list('id_lenh_san_xuat__id_lenh_san_xuat', flat=True).distinct()
        
        return JsonResponse({
            'lenh_sx_list': list(lenh_sx_with_wo)
        })
    except Exception as e:
        return JsonResponse({
            'lenh_sx_list': [],
            'error': str(e)
        }, status=400)


def get_lenh_sx_by_don_hang_wo(request, ma_don_hang):
    """API endpoint to get production orders by order ID"""
    try:
        lenh_sx_list = LenhSanXuat.objects.filter(
            id_don_hang=ma_don_hang
        ).values_list('id_lenh_san_xuat', flat=True)
        
        return JsonResponse({
            'status': 'success',
            'data': list(lenh_sx_list)
        })
    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': f'Lỗi: {str(e)}'
        })


def get_purchase_records_by_lenh_sx_wo(request, ma_lenh_sx):
    """API endpoint to get available purchase records by production order"""
    try:
        # Sửa: Lấy purchase records chưa có WO
        existing_wo_purchase_ids = BangKeWo.objects.values_list('id_bang_ke_thu_mua_tu_dan', flat=True)
        
        purchase_records = BangKeThuMuaTuDan.objects.filter(
            id_lenh_san_xuat__id_lenh_san_xuat=ma_lenh_sx
        ).exclude(
            id_bang_ke_thu_mua_tu_dan__in=existing_wo_purchase_ids
        ).select_related('id_lenh_san_xuat')
        
        formatted_records = []
        for record in purchase_records:
            # Get material info
            try:
                vat_tu = VatTu.objects.get(id_san_pham=record.id_san_pham)
                ten_nguyen_lieu = vat_tu.ten_khac or ''
                ma_hs = vat_tu.ma_hs or ''
                don_vi_tinh = vat_tu.don_vi_tinh or 'KGM'
            except VatTu.DoesNotExist:
                ten_nguyen_lieu = f"Sản phẩm {record.id_san_pham}"
                ma_hs = ''
                don_vi_tinh = 'KGM'
            
            # Get rollback quantity
            rollback_record = BangKeTruLuiNguyenLieu.objects.filter(
                id_lenh_san_xuat=record.id_lenh_san_xuat,
                id_san_pham=record.id_san_pham
            ).first()
            so_luong_san_pham_xuat = rollback_record.so_luong_san_pham_xuat if rollback_record else 0
            
            formatted_records.append({
                'id_bang_ke_thu_mua_tu_dan': record.id_bang_ke_thu_mua_tu_dan,
                'ma_lenh_sx': record.id_lenh_san_xuat.id_lenh_san_xuat,
                'ma_don_hang': record.id_lenh_san_xuat.id_don_hang,
                'ten_nguyen_lieu': ten_nguyen_lieu,
                'ma_hs': ma_hs,
                'don_vi_tinh': don_vi_tinh,
                'so_luong_san_pham_xuat': so_luong_san_pham_xuat,
                #'ngay_lap_giay_to': record.ngay_lap_giay_to.strftime('%d/%m/%Y') if record.ngay_lap_giay_to else '',
            })
        
        return JsonResponse({
            'status': 'success',
            'purchase_records': formatted_records
        })
    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': f'Lỗi khi tải danh sách bảng kê thu mua: {str(e)}'
        }, status=400)


def get_purchase_details_wo(request, purchase_id):
    """API endpoint to get detailed purchase information for table display"""
    try:
        # Get purchase record
        purchase_record = get_object_or_404(BangKeThuMuaTuDan, pk=purchase_id)
        
        # Get detailed purchase records (chi tiết thu mua)
        detail_records = CtBangKeThuMuaTuDan.objects.filter(
            id_bang_ke_thu_mua_tu_dan=purchase_record
        ).select_related('id_nguoi_ban')
        
        details = []
        for detail in detail_records:
            nguoi = detail.id_nguoi_ban
            
            # Format CCCD với ngày cấp
            cccd_info = ''
            if nguoi and nguoi.so_cmnd_cccd:
                cccd_info = nguoi.so_cmnd_cccd
                if nguoi.ngay_cap_cmnd_cccd:
                    cccd_info += f" cấp ngày {nguoi.ngay_cap_cmnd_cccd.strftime('%d/%m/%Y')}"
            
            details.append({
                'ngay_mua_hang': detail.ngay_mua_hang.strftime('%Y-%m-%d') if detail.ngay_mua_hang else '',
                'ten_nguoi_ban': nguoi.ten if nguoi else '',
                'dia_chi': nguoi.dia_chi if nguoi else '',
                'so_cmnd_cccd': cccd_info,  # Đã format với ngày cấp
                'so_luong': detail.so_luong or 0,
                'don_gia': detail.don_gia or 0,
                'ghi_chu': detail.ghi_chu or '',
            })
        
        return JsonResponse({
            'status': 'success',
            'details': details
        })
    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': f'Lỗi khi tải chi tiết bảng kê thu mua: {str(e)}'
        }, status=400)
    
@require_GET
def wo_export_pdf(request, pk):
    """Render WO record as HTML for PDF printing (AJAX)"""
    try:
        wo_record = BangKeWo.objects.select_related(
            'id_lenh_san_xuat', 
            'id_nguoi'
        ).get(pk=pk)
    except BangKeWo.DoesNotExist:
        return HttpResponse("Không tìm thấy bảng kê WO", status=404)

    # Get material info từ VatTu
    try:
        vat_tu = VatTu.objects.get(id_san_pham=wo_record.id_san_pham)
        ten_nguyen_lieu = vat_tu.ten_khac or vat_tu.ten_sp_chinh
        ma_hs = vat_tu.ma_hs or ''
        don_vi_tinh = vat_tu.don_vi_tinh or 'KGM'
    except VatTu.DoesNotExist:
        ten_nguyen_lieu = f"Sản phẩm {wo_record.id_san_pham}"
        ma_hs = ''
        don_vi_tinh = 'KGM'
    # Get purchase details for table display
    purchase_details = []
    total_quantity = 0
    total_amount = 0
    
    try:
        detail_records = CtBangKeThuMuaTuDan.objects.filter(
            id_bang_ke_thu_mua_tu_dan=wo_record.id_bang_ke_thu_mua_tu_dan
        ).select_related('id_nguoi_ban').order_by('ngay_mua_hang')
        
        for detail in detail_records:
            nguoi = detail.id_nguoi_ban
            
            # Format CCCD với ngày cấp
            cccd_info = ''
            if nguoi and nguoi.so_cmnd_cccd:
                cccd_info = nguoi.so_cmnd_cccd
                if nguoi.ngay_cap_cmnd_cccd:
                    cccd_info += f" cấp ngày {nguoi.ngay_cap_cmnd_cccd.strftime('%d/%m/%Y')}"
            
            thanh_tien = (detail.so_luong or 0) * (detail.don_gia or 0)
            total_quantity += (detail.so_luong or 0)
            total_amount += thanh_tien
            
            purchase_details.append({
                'ngay_mua_hang': detail.ngay_mua_hang.strftime('%d/%m/%Y') if detail.ngay_mua_hang else '',
                'ten_nguoi_ban': nguoi.ten if nguoi else '',
                'dia_chi': nguoi.dia_chi if nguoi else '',
                'so_cmnd_cccd': cccd_info,
                'so_luong': detail.so_luong or 0,
                'don_gia': detail.don_gia or 0,
                'thanh_tien': thanh_tien,
                'ghi_chu': detail.ghi_chu or '',
            })
    except Exception as e:
        print(f"Error loading purchase details: {e}")

    # Format WO record data
    formatted_record = {
        'ten_thuong_nhan': 'Công ty cổ phần Tân Phong',
        'ma_so_thue': '2600274542',
        'ten_nguyen_lieu': wo_record.ten_hang_hoa or '',  
        'ten_hang_hoa': ten_nguyen_lieu,
        'ma_hs': ma_hs,
        'so_luong_wo': wo_record.so_luong or 0,
        'don_vi_tinh': don_vi_tinh,
        'tri_gia_fob': wo_record.tri_gia_fob or 0,
        'to_khai_hai_quan': wo_record.to_khai_hai_quan or '',
        'dia_chi_thu_mua': wo_record.dia_chi_thu_mua or '',
        'noi_khai_thac': wo_record.noi_khai_thac or '',
        'ten_nguoi': wo_record.id_nguoi.ten if wo_record.id_nguoi else '',
        'cccd_cmnd': wo_record.id_nguoi.so_cmnd_cccd if wo_record.id_nguoi else '',
        'ngay_display': wo_record.ngay.strftime('%d tháng %m năm %Y') if wo_record.ngay else datetime.now().strftime('%d tháng %m năm %Y'),
        'ma_lenh_sx': wo_record.id_lenh_san_xuat.id_lenh_san_xuat,
        'ma_don_hang': wo_record.id_lenh_san_xuat.id_don_hang,
    }
    
    context = {
        'record': formatted_record,
        'purchase_details': purchase_details,
        'total_quantity': total_quantity,
        'total_amount': total_amount,
    }

    return render(request, 'wo_export_pdf.html', context)

def wo_export_word(request, pk):
    """Export WO record to Word document"""
    try:   
        wo_record = BangKeWo.objects.select_related(
            'id_lenh_san_xuat', 
            'id_nguoi'
        ).get(pk=pk)
    except BangKeWo.DoesNotExist:
        return HttpResponse("Không tìm thấy bảng kê WO", status=404)
    except ImportError:
        return HttpResponse("Thiếu thư viện python-docx. Vui lòng cài đặt: pip install python-docx", status=500)

    # Get material info từ VatTu
    try:
        vat_tu = VatTu.objects.get(id_san_pham=wo_record.id_san_pham)
        ten_nguyen_lieu = vat_tu.ten_khac or vat_tu.ten_sp_chinh
        ma_hs = vat_tu.ma_hs or ''
        don_vi_tinh = vat_tu.don_vi_tinh or 'KGM'
    except VatTu.DoesNotExist:
        ten_nguyen_lieu = f"Sản phẩm {wo_record.id_san_pham}"
        ma_hs = ''
        don_vi_tinh = 'KGM'

    # Get purchase details
    purchase_details = []
    total_quantity = 0
    total_amount = 0
    
    try:
        detail_records = CtBangKeThuMuaTuDan.objects.filter(
            id_bang_ke_thu_mua_tu_dan=wo_record.id_bang_ke_thu_mua_tu_dan
        ).select_related('id_nguoi_ban').order_by('ngay_mua_hang')
        
        for detail in detail_records:
            nguoi = detail.id_nguoi_ban
            
            # Format CCCD với ngày cấp
            cccd_info = ''
            if nguoi and nguoi.so_cmnd_cccd:
                cccd_info = nguoi.so_cmnd_cccd
                if nguoi.ngay_cap_cmnd_cccd:
                    cccd_info += f" cấp ngày {nguoi.ngay_cap_cmnd_cccd.strftime('%d/%m/%Y')}"
            
            thanh_tien = (detail.so_luong or 0) * (detail.don_gia or 0)
            total_quantity += (detail.so_luong or 0)
            total_amount += thanh_tien
            
            purchase_details.append({
                'ngay_mua_hang': detail.ngay_mua_hang.strftime('%d/%m/%Y') if detail.ngay_mua_hang else '',
                'ten_nguoi_ban': nguoi.ten if nguoi else '',
                'dia_chi': nguoi.dia_chi if nguoi else '',
                'so_cmnd_cccd': cccd_info,
                'so_luong': detail.so_luong or 0,
                'don_gia': detail.don_gia or 0,
                'thanh_tien': thanh_tien,
                'ghi_chu': detail.ghi_chu or '',
            })
    except Exception as e:
        print(f"Error loading purchase details: {e}")

    # Create Word document
    doc = Document()
    
    # Set document margins to A4 Landscape (khổ ngang)
    sections = doc.sections
    for section in sections:
        section.orientation = WD_ORIENT.LANDSCAPE  # Khổ ngang
        section.page_width = Inches(11.69) # A4 landscape width
        section.page_height = Inches(8.27)  # A4 landscape height
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # Header - Phụ lục II (căn giữa)
    header1 = doc.add_paragraph()
    header1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header1.space_after = Pt(0)
    run1 = header1.add_run('Phụ lục II')
    run1.font.name = 'Times New Roman'
    run1.font.size = Pt(12)  # Giảm từ 13 xuống 12
    run1.bold = True

    # Main title (căn giữa)
    header2 = doc.add_paragraph()
    header2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header2.space_after = Pt(0)
    run2 = header2.add_run('BẢNG KÊ KHAI HÀNG HÓA XUẤT KHẨU ĐẠT TIÊU CHÍ "WO"')
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(14)
    run2.bold = True

    # Subtitle (căn giữa, in đậm, cỡ 14)
    sub_para1 = doc.add_paragraph()
    sub_para1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_para1.space_after = Pt(0)
    run_sub1 = sub_para1.add_run('(sử dụng nguyên liệu thu mua trong nước, không có hóa đơn giá trị gia tăng)')
    run_sub1.font.name = 'Times New Roman'
    run_sub1.font.size = Pt(14)  # Tăng từ 10 lên 14
    run_sub1.bold = True  # Thêm in đậm, bỏ in nghiêng

    # Subtitles tiếp theo
    sub_para2 = doc.add_paragraph()
    sub_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_para2.space_after = Pt(0)
    run_sub2 = sub_para2.add_run('(ban hành kèm theo Thông tư số 74/2023/TT-BCT ngày 29/12/2023')
    run_sub2.font.name = 'Times New Roman'
    run_sub2.font.size = Pt(13)
    run_sub2.italic = True

    sub_para3 = doc.add_paragraph()
    sub_para3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_para3.space_after = Pt(0)
    run_sub3 = sub_para3.add_run('Của Bộ trưởng Bộ Công Thương sửa đổi, bổ sung một số điều của Thông tư 05/2018/TT-BCT)')
    run_sub3.font.name = 'Times New Roman'
    run_sub3.font.size = Pt(13)
    run_sub3.italic = True

    # Company info - Sử dụng bảng 3 cột như trong mẫu, bỏ row 6
    info_table = doc.add_table(rows=5, cols=3)  # Giảm từ 6 xuống 5 rows
    info_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    
    # Xóa borders của bảng info
    def remove_table_borders(table):
        tbl = table._tbl
        for row in tbl.tr_lst:
            for cell in row.tc_lst:
                tcPr = cell.tcPr
                tcBorders = OxmlElement('w:tcBorders')
                for border_name in ['top', 'left', 'bottom', 'right']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'nil')
                    tcBorders.append(border)
                tcPr.append(tcBorders)
    
    remove_table_borders(info_table)
    
    # Set column widths
    info_table.columns[0].width = Inches(3.0)
    info_table.columns[1].width = Inches(2.5) 
    info_table.columns[2].width = Inches(2.2)

    # Fill company info theo layout mẫu - bỏ row 6
    # Row 1
    info_table.cell(0, 0).text = 'Tên Thương nhân: Công ty cổ phần Tân Phong'
    info_table.cell(0, 1).text = ''
    info_table.cell(0, 2).text = 'Tiêu chí áp dụng: WO'
    
    # Row 2  
    info_table.cell(1, 0).text = f'Mã số thuế: 2600274542'
    info_table.cell(1, 1).text = ''
    info_table.cell(1, 2).text = f'Tên hàng hóa: {ten_nguyen_lieu or ""}'
    
    # Row 3
    info_table.cell(2, 0).text = f'Tờ khai hải quan xuất khẩu số: {wo_record.to_khai_hai_quan or ""}'
    info_table.cell(2, 1).text = ''
    info_table.cell(2, 2).text = f'Mã HS của hàng hóa (6 số): {ma_hs}'
    
    # Row 4
    info_table.cell(3, 0).text = f'Địa chỉ nơi tổ chức thu mua: {wo_record.dia_chi_thu_mua or ""}'
    info_table.cell(3, 1).text = ''
    info_table.cell(3, 2).text = f'Số lượng: {wo_record.so_luong or 0:.0f} đơn vị tính KGM'
    
    # Row 5
    nguoi_phu_trach = f'{wo_record.id_nguoi.ten if wo_record.id_nguoi else ""}, CCCD số: {wo_record.id_nguoi.so_cmnd_cccd if wo_record.id_nguoi else ""}'
    info_table.cell(4, 0).text = f'Người phụ trách thu mua (Tên, số định danh cá nhân): {nguoi_phu_trach}'
    info_table.cell(4, 1).text = ''
    info_table.cell(4, 2).text = f'Trị giá (FOB): {wo_record.tri_gia_fob or 0:.1f} USD'

    # Format info table text
    for row in info_table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.space_after = Pt(2)  # Giảm khoảng cách
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)

    # Minimal space before main table
    space_before_table = doc.add_paragraph()
    space_before_table.space_after = Pt(3)

    # Main data table - Tạo table với cấu trúc đúng như mẫu
    table = doc.add_table(rows=4 + len(purchase_details), cols=11)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Set specific column widths to prevent text wrapping
    widths = [Inches(0.8), Inches(1.0), Inches(1.2), Inches(1.3), Inches(1.0), 
              Inches(0.6), Inches(1.0), Inches(0.8), Inches(0.9), Inches(1.0), Inches(0.8)]
    for i, width in enumerate(widths):
        table.columns[i].width = width
    
    # Set table borders
    def set_table_borders(table):
        tbl = table._tbl
        for row in tbl.tr_lst:
            for cell in row.tc_lst:
                tcPr = cell.tcPr
                tcBorders = OxmlElement('w:tcBorders')
                
                for border_name in ['top', 'left', 'bottom', 'right']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'single')
                    border.set(qn('w:sz'), '4')
                    border.set(qn('w:space'), '0')
                    border.set(qn('w:color'), '000000')
                    tcBorders.append(border)
                
                tcPr.append(tcBorders)
    
    set_table_borders(table)
    
    # Header row 1 - Merge cells như mẫu
    table.cell(0, 0).text = 'Ngày tháng năm mua hàng'
    
    # Merge cells cho "Người bán" (cột 1-3)
    table.cell(0, 1).merge(table.cell(0, 3))
    table.cell(0, 1).text = 'Người bán'
    
    # Merge cells cho "Nguyên liệu thu mua trong nước" (cột 4-9)
    table.cell(0, 4).merge(table.cell(0, 9))
    table.cell(0, 4).text = 'Nguyên liệu thu mua trong nước'
    
    table.cell(0, 10).text = 'Ghi chú'

    # Header row 2 - Chi tiết cột
    table.cell(1, 0).text = ''
    table.cell(1, 1).text = 'Tên người bán'
    table.cell(1, 2).text = 'Địa chỉ'
    table.cell(1, 3).text = 'Số định danh cá nhân (số CCCD) và ngày cấp'
    table.cell(1, 4).text = 'Tên nguyên liệu'
    table.cell(1, 5).text = 'Mã HS'
    table.cell(1, 6).text = 'Nơi khai thác/đánh bắt/nuôi trồng'
    table.cell(1, 7).text = 'Số lượng và Đơn vị tính (kg)'
    table.cell(1, 8).text = 'Đơn giá (VND)'
    table.cell(1, 9).text = 'Tổng trị giá (VND)'
    table.cell(1, 10).text = ''

    # Header row 3 - Số thứ tự cột
    column_numbers = ['(1)', '(2)', '(3)', '(4)', '(5)', '(6)', '(7)', '(8)', '(9)', '(10)', '(11)']
    for i, num in enumerate(column_numbers):
        table.cell(2, i).text = num

    # Format headers
    for i in range(3):
        for j in range(11):
            try:
                cell = table.cell(i, j)
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(10)
                        if i < 2:
                            run.bold = True
                        else:
                            run.italic = True
            except:
                continue

    # Fill data rows
    for idx, detail in enumerate(purchase_details):
        row_idx = 3 + idx
        
        data_values = [
            detail['ngay_mua_hang'],
            detail['ten_nguoi_ban'],
            detail['dia_chi'],
            detail['so_cmnd_cccd'],
            wo_record.ten_hang_hoa or ten_nguyen_lieu,
            ma_hs,
            wo_record.noi_khai_thac or '',
            f"{detail['so_luong']:.0f}",
            f"{detail['don_gia']:,.0f}",
            f"{detail['thanh_tien']:,.0f}",
            detail['ghi_chu']
        ]
        
        for col_idx, value in enumerate(data_values):
            cell = table.cell(row_idx, col_idx)
            cell.text = str(value)
            
            # Format data cells
            for paragraph in cell.paragraphs:
                if col_idx in [0, 3]:  # Ngày và CCCD center
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif col_idx in [7, 8, 9]:  # Số liệu right align
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:  # Còn lại left align
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)

    # Total row
    total_row_idx = 3 + len(purchase_details)
    
    # Merge cells cho "Tổng cộng"
    table.cell(total_row_idx, 0).merge(table.cell(total_row_idx, 6))
    table.cell(total_row_idx, 0).text = 'Tổng cộng'
    
    table.cell(total_row_idx, 7).text = f"{total_quantity:.0f} KGM"
    table.cell(total_row_idx, 8).text = ''
    table.cell(total_row_idx, 9).text = f"{total_amount:,.0f}"
    table.cell(total_row_idx, 10).text = ''
    
    # Format total row
    for col_idx in [0, 7, 9]:
        cell = table.cell(total_row_idx, col_idx)
        for paragraph in cell.paragraphs:
            if col_idx == 0:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                run.bold = True

    # Conclusion - Remove spacing after
    conclusion_para = doc.add_paragraph()
    conclusion_para.space_after = Pt(0)  # Remove spacing
    conclusion_run = conclusion_para.add_run('Kết luận: Hàng hóa đáp ứng tiêu chí "WO"')
    conclusion_run.font.name = 'Times New Roman'
    conclusion_run.font.size = Pt(10)
    conclusion_run.bold = True

    commitment_para = doc.add_paragraph()
    commitment_para.space_after = Pt(0)  # Remove spacing
    commitment_run = commitment_para.add_run('Công ty cam kết số liệu, thông tin khai báo trên là đúng và chịu trách nhiệm trước pháp luật về thông tin, số liệu đã khai.')
    commitment_run.font.name = 'Times New Roman'
    commitment_run.font.size = Pt(10)

    # Signature section - Căn theo vị trí cột số lượng/đơn giá với độ rộng chuẩn
    ngay_display = wo_record.ngay.strftime('Ngày %d tháng %m năm %Y') if wo_record.ngay else datetime.now().strftime('Ngày %d tháng %m năm %Y')
    
    # Tạo bảng để căn chỉnh chữ ký theo đúng vị trí cột 7-9 (số lượng, đơn giá, tổng trị giá)
    signature_table = doc.add_table(rows=1, cols=11)
    signature_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Set same column widths as main table
    for i, width in enumerate(widths):
        signature_table.columns[i].width = width
    
    # Xóa borders của bảng signature
    def remove_signature_table_borders(table):
        tbl = table._tbl
        for row in tbl.tr_lst:
            for cell in row.tc_lst:
                tcPr = cell.tcPr
                tcBorders = OxmlElement('w:tcBorders')
                for border_name in ['top', 'left', 'bottom', 'right']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'nil')
                    tcBorders.append(border)
                tcPr.append(tcBorders)
    
    remove_signature_table_borders(signature_table)
    
    # Merge cột 7-9 (vị trí số lượng, đơn giá, tổng trị giá) để đặt chữ ký
    signature_table.cell(0, 7).merge(signature_table.cell(0, 9))
    
    # Đặt nội dung ngày vào cột đã merge
    signature_table.cell(0, 7).text = ngay_display
    
    # Format ngày
    cell = signature_table.cell(0, 7)
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.space_after = Pt(0)  # Remove spacing
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run.italic = True

    # Tạo bảng thứ 2 cho chức danh
    signature_table2 = doc.add_table(rows=1, cols=11)
    signature_table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Set same column widths as main table
    for i, width in enumerate(widths):
        signature_table2.columns[i].width = width
    
    remove_signature_table_borders(signature_table2)
    
    # Merge cột 7-9 cho chức danh
    signature_table2.cell(0, 7).merge(signature_table2.cell(0, 9))
    signature_table2.cell(0, 7).text = 'Người đại diện theo pháp luật của thương nhân'
    
    # Format chức danh
    cell = signature_table2.cell(0, 7)
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.space_after = Pt(0)  # Remove spacing
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run.bold = True

    # Tạo bảng thứ 3 cho ghi chú ký
    signature_table3 = doc.add_table(rows=1, cols=11)
    signature_table3.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Set same column widths as main table
    for i, width in enumerate(widths):
        signature_table3.columns[i].width = width
    
    remove_signature_table_borders(signature_table3)
    
    # Merge cột 7-9 cho ghi chú ký
    signature_table3.cell(0, 7).merge(signature_table3.cell(0, 9))
    signature_table3.cell(0, 7).text = '(Ký, đóng dấu, ghi rõ họ, tên)'
    
    # Format ghi chú ký
    cell = signature_table3.cell(0, 7)
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.space_after = Pt(0)  # Remove spacing
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run.italic = True

    # Space for signature - giảm spacing
    for i in range(4):
        signature_space = doc.add_paragraph()
        signature_space.space_after = Pt(0)  # Remove spacing

    # Notes section - Remove spacing
    note_title = doc.add_paragraph()
    note_title.space_after = Pt(0)  # Remove spacing
    note_title_run = note_title.add_run('Ghi chú:')
    note_title_run.font.name = 'Times New Roman'
    note_title_run.font.size = Pt(10)
    note_title_run.bold = True

    note1 = doc.add_paragraph()
    note1.space_after = Pt(0)  # Remove spacing
    note1_run = note1.add_run('- Mẫu Bảng kê khai này áp dụng trong trường hợp nguyên liệu được thu mua trong nước để sản xuất ra hàng hóa xuất khẩu nhưng không có hóa đơn giá trị gia tăng.')
    note1_run.font.name = 'Times New Roman'
    note1_run.font.size = Pt(10)

    note2 = doc.add_paragraph()
    note2.space_after = Pt(0)  # Remove spacing
    note2_run = note2.add_run('- Thương nhân nộp bản sao các chứng từ (đóng dấu sao y bản chính): Quy trình sản xuất hàng hóa, Giấy CCCD của người bán nguyên liệu; Giấy xác nhận của người bán nguyên liệu về vùng nuôi trồng, khai thác, số lượng và trị giá bán cho thương nhân (nếu có) để đối chiếu với thông tin kê khai.')
    note2_run.font.name = 'Times New Roman'
    note2_run.font.size = Pt(10)

    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # Create response
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    
    filename = f'Bang_ke_WO_{wo_record.id_lenh_san_xuat.id_lenh_san_xuat}_{wo_record.id}.docx'
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    
    return response